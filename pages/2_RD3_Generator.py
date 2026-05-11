import streamlit as st
import pandas as pd
from io import BytesIO
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── PAGE CONFIG ────────────────────────────────────────────────────────────
st.set_page_config(page_title="RD3 Generator | SOCOTEC Arabia", page_icon="🏗️", layout="wide")

st.markdown("""
<style>
    [data-testid="stSidebar"] { min-width:220px; max-width:220px; }
    .main-header {
        background: linear-gradient(135deg, #0072BB, #005a94);
        padding: 20px 28px; border-radius: 8px; margin-bottom: 20px;
    }
    .main-header h2 { color:white; margin:0; font-size:1.5rem; }
    .main-header p  { color:rgba(255,255,255,0.82); margin:4px 0 0 0; font-size:0.88rem; }
    .section-title  { font-weight:700; color:#0072BB; font-size:1rem; margin:16px 0 8px 0;
                      border-bottom:2px solid #e0eaf3; padding-bottom:4px; }
    .help-note      { background:#f0f7ff; border-left:3px solid #0072BB;
                      padding:8px 12px; border-radius:4px; font-size:0.85rem; color:#444; }
    div[data-testid="stTabs"] button { font-size:0.85rem; }
    .stDownloadButton > button { background:#0072BB!important; color:white!important;
                                  font-weight:600!important; border-radius:6px!important; }
</style>
""", unsafe_allow_html=True)

st.markdown("""
<div class="main-header">
    <h2>🏗️ RD3 Generator — Waterproofness Final Report</h2>
    <p>Fill in all sections, then click <strong>Generate Report</strong> at the bottom to download the completed .docx</p>
</div>
""", unsafe_allow_html=True)

# ─── HELPER: section title ───────────────────────────────────────────────────
def sec(label):
    st.markdown(f'<div class="section-title">{label}</div>', unsafe_allow_html=True)

# ─── FORM TABS ───────────────────────────────────────────────────────────────
tab_info, tab_proj, tab_works, tab_dates, tab_conclusion, tab_a1, tab_a2, tab_a3 = st.tabs([
    "1 · Report Info", "2 · Project", "3 · WP Works & Visits",
    "4 · Reservations", "5 · Conclusion",
    "Annex 1 – Roofs", "Annex 2 – Façade", "Annex 3 – Basement"
])

# ═══════════════════════════════════════════════════════════════════════════════
# TAB 1 – REPORT INFO
# ═══════════════════════════════════════════════════════════════════════════════
with tab_info:
    sec("Document Reference")
    c1, c2 = st.columns(2)
    ref_num         = c1.text_input("Reference Number", placeholder="e.g. SOCOTEC-RD3-2026-001")
    tawuniya_id     = c2.text_input("Tawuniya Visit ID", placeholder="e.g. TW-12345")
    c3, c4, c5 = st.columns(3)
    doc_ref         = c3.text_input("Document Reference", placeholder="e.g. RD3/KSA/001")
    tis_agency      = c4.text_input("TIS Agency", value="SOCOTEC Arabia")
    version         = c5.text_input("Version", value="1.0")
    issue_date      = st.date_input("Date of Issue", value=date.today())

    sec("Responsible Expert")
    c6, c7, c8, c9 = st.columns(4)
    expert_name     = c6.text_input("Surname, Name")
    expert_phase    = c7.text_input("Phase", placeholder="e.g. Final")
    expert_degree   = c8.text_input("Degree", placeholder="e.g. Civil Eng.")
    expert_spec     = c9.text_input("Speciality", placeholder="e.g. Structures")

    sec("Author of This Report")
    c10, c11, c12 = st.columns(3)
    author_name     = c10.text_input("Author Name")
    author_phone    = c11.text_input("Phone Number", placeholder="+966 5X XXX XXXX")
    author_email    = c12.text_input("Email", placeholder="name@socotec.com")
    city            = st.text_input("City (for signature block)", value="Riyadh")

# ═══════════════════════════════════════════════════════════════════════════════
# TAB 2 – PROJECT INFO
# ═══════════════════════════════════════════════════════════════════════════════
with tab_proj:
    sec("I – Project Information")
    proj_title      = st.text_input("Project Title / Name")
    proj_address    = st.text_area("Address of the Premises", height=80)
    c1, c2 = st.columns(2)
    ref_rd0         = c1.text_input("Reference RD0", placeholder="e.g. RD0/2025/001")
    owner           = c2.text_input("Principal / Owner")
    buildings_use   = st.text_area("Buildings Included in the Project and Their Use", height=100,
                                    placeholder="e.g. Building A – Residential (20 floors), Building B – Commercial (5 floors)")

# ═══════════════════════════════════════════════════════════════════════════════
# TAB 3 – WP WORKS & DATES
# ═══════════════════════════════════════════════════════════════════════════════
with tab_works:
    sec("II – Waterproofing Works Concerned")
    st.markdown('<div class="help-note">Select which WP works are covered. Each checked item will enable its Annex tab.</div>', unsafe_allow_html=True)
    cc1, cc2, cc3 = st.columns(3)
    wp_roofs        = cc1.checkbox("🏠 Roofs (→ Annex 1)",    value=False)
    wp_facades      = cc2.checkbox("🏛️ Façades (→ Annex 2)", value=False)
    wp_basements    = cc3.checkbox("🔩 Basements (→ Annex 3)", value=False)

    sec("III.1 – Occupancy Certificate")
    occ_date_str    = st.text_input("Date of Occupancy Certificate", placeholder="dd/mm/yyyy or TBD")
    occ_status      = st.radio("Status", ["Expected", "Confirmed"], horizontal=True)

    sec("III.2 – WP Control Site Visits")
    st.markdown("Add all site visits for waterproofing control:")

    default_visits = pd.DataFrame({
        "Reference":        [""],
        "Date":             [""],
        "Inspector":        [""],
        "Part Inspected":   [""]
    })
    visits_df = st.data_editor(
        default_visits, num_rows="dynamic", use_container_width=True,
        column_config={
            "Reference":      st.column_config.TextColumn("Reference", width="medium"),
            "Date":           st.column_config.TextColumn("Date", width="small"),
            "Inspector":      st.column_config.TextColumn("Inspector", width="medium"),
            "Part Inspected": st.column_config.TextColumn("Part Inspected", width="large"),
        }, key="visits_table"
    )

    sec("Defects / Disorders Not Repaired")
    defects_note    = st.text_area("List any defects/disorders not repaired at date of this report (not being a technical reservation)",
                                    height=100, placeholder="None observed / describe defects...")

# ═══════════════════════════════════════════════════════════════════════════════
# TAB 4 – TECHNICAL RESERVATIONS
# ═══════════════════════════════════════════════════════════════════════════════
with tab_dates:
    sec("IV – List of Technical Reservations Not Closed")
    reserv_exist    = st.radio("Technical reservations issued and not closed at date of this report?",
                               ["YES", "NO"], index=1, horizontal=True, key="main_reserv")
    default_reserv  = pd.DataFrame({
        "Concept":              [""],
        "Issuing Cert (ref/date)": [""],
        "Expected Closing / FINAL": [""]
    })
    reserv_df = st.data_editor(
        default_reserv, num_rows="dynamic", use_container_width=True,
        column_config={
            "Concept":              st.column_config.TextColumn("Concept", width="large"),
            "Issuing Cert (ref/date)": st.column_config.TextColumn("Issuing Cert (ref/date)", width="medium"),
            "Expected Closing / FINAL": st.column_config.TextColumn("Expected Closing / FINAL", width="medium"),
        }, key="main_reserv_table"
    )

# ═══════════════════════════════════════════════════════════════════════════════
# TAB 5 – FINAL CONCLUSION (MAIN)
# ═══════════════════════════════════════════════════════════════════════════════
with tab_conclusion:
    sec("V – Final Conclusion (Design + Execution)")
    main_conclusion = st.radio(
        "Is the execution of works adapted to the project? Including eventual modifications?",
        ["YES", "NO"], horizontal=True, key="main_concl"
    )
    main_concl_text = st.text_area("Please expose and develop your answer in detail:", height=200,
                                    placeholder="Describe the overall waterproofing assessment conclusion...")

# ═══════════════════════════════════════════════════════════════════════════════
# ANNEX 1 – ROOFS
# ═══════════════════════════════════════════════════════════════════════════════
with tab_a1:
    if not wp_roofs:
        st.info("Enable **Roofs** in Tab 3 – WP Works to unlock this annex.")
    else:
        sec("Annex 1 – Waterproofing of Roofs")
        st.markdown("**Type of Roofs:**")
        ra1, ra2, ra3, ra4 = st.columns(4)
        a1_roof         = ra1.checkbox("Roof")
        a1_terrace_top  = ra2.checkbox("Rooftop Terrace")
        a1_terrace_int  = ra3.checkbox("Intermediate Terrace")
        a1_patios       = ra4.checkbox("Patios")
        a1_other_type   = st.text_input("Other roof type:", key="a1_other_type")

        sec("I – Works Description")
        a1_desc1 = st.text_area("I.1. Describe the roof concerned (type, materials, layers, slope, location):", height=90, key="a1d1")
        a1_desc2 = st.text_area("I.2. Describe the WP system layers (material, manufacturer, thickness) from inside outwards:", height=90, key="a1d2")
        a1_desc3 = st.text_area("I.3. Describe the junctions of the WP system with other elements:", height=80, key="a1d3")
        a1_innovative = st.radio("I.4. Does the WP system include innovative technique, procedure, or materials?", ["YES","NO"], index=1, horizontal=True, key="a1inn")
        if a1_innovative == "YES":
            a1_innovative_desc = st.text_area("Describe (including RD2 reference):", height=80, key="a1inndesc")
        else:
            a1_innovative_desc = ""

        sec("II – Materials of the Waterproofing System")
        st.markdown("**Delivery Orders & Material Certificates**")
        ca1, ca2 = st.columns(2)
        a1_do_avail     = ca1.radio("Delivery orders and/or certificates available?", ["YES","NO"], horizontal=True, key="a1do")
        a1_do_compliant = ca2.radio("WP materials used compliant with design?", ["YES","NO"], horizontal=True, key="a1comp")
        a1_do_docs      = st.text_area("Reviewed documents (delivery orders):", height=70, placeholder="1.\n2.\n3.", key="a1dodocs")
        st.markdown("**Tests and Quality Control Reports**")
        a1_ponding      = st.radio("Water ponding test report available?", ["YES","NO"], horizontal=True, key="a1pond")
        a1_test_docs    = st.text_area("Reviewed documents (tests):", height=70, placeholder="1.\n2.\n3.", key="a1testdocs")
        a1_other_docs   = st.text_area("Other reviewed documents:", height=70, placeholder="1.\n2.", key="a1otherdocs")

        sec("III – Modifications Regarding RD0")
        cb1, cb2 = st.columns(2)
        a1_mod          = cb1.radio("Any modification of WP system after RD0?", ["YES","NO"], index=1, horizontal=True, key="a1mod")
        a1_mod_valid    = cb2.radio("If YES — was modification validated (design/calculi)?", ["YES","NO","N/A"], index=2, horizontal=True, key="a1modval")
        cc1b, cc2b = st.columns(2)
        a1_tis_accept   = cc1b.radio("Does TIS find modification acceptable?", ["YES","NO","N/A"], index=2, horizontal=True, key="a1tis")
        a1_exec_ok      = cc2b.radio("Is the execution satisfactory?", ["YES","NO"], horizontal=True, key="a1exec")
        a1_reserv_refs  = st.text_input("Technical Reservation reference(s) if NO:", key="a1refs")

        sec("IV – Roof Waterproofing Tests")
        a1_defects      = st.radio("Did you observe any defect or sign of damage regarding water roof infiltration?", ["YES","NO"], horizontal=True, key="a1def")
        a1_defect_ref   = st.text_input("If YES — Technical Reservation ref:", key="a1defref")

        sec("V – Technical Reservations (Annex 1)")
        st.markdown("**Reservations Issued and CLOSED:**")
        a1_closed_exist = st.radio("Any closed technical reservations?", ["YES","NO"], index=1, horizontal=True, key="a1cl")
        a1_closed_df    = st.data_editor(pd.DataFrame({"Concept":[""],"Issuing Cert (ref/date)":[""],"Closing Cert (ref/date)":[""]}),
                            num_rows="dynamic", use_container_width=True, key="a1cldf")
        st.markdown("**Reservations Issued and NOT CLOSED:**")
        a1_open_exist   = st.radio("Any open technical reservations?", ["YES","NO"], index=1, horizontal=True, key="a1op")
        a1_open_df      = st.data_editor(pd.DataFrame({"Concept":[""],"Issuing Cert (ref/date)":[""],"Expected Closing / FINAL":[""]}),
                            num_rows="dynamic", use_container_width=True, key="a1opdf")

        sec("VI – Final Conclusion (Annex 1)")
        a1_conclusion   = st.radio("Is execution adapted to project including modifications?", ["YES","NO"], horizontal=True, key="a1concl")
        a1_concl_text   = st.text_area("Develop your answer:", height=150, key="a1concltxt")

# ═══════════════════════════════════════════════════════════════════════════════
# ANNEX 2 – FAÇADE
# ═══════════════════════════════════════════════════════════════════════════════
with tab_a2:
    if not wp_facades:
        st.info("Enable **Façades** in Tab 3 – WP Works to unlock this annex.")
    else:
        sec("Annex 2 – Waterproofing of Façade")
        st.markdown("**Type of Façade:**")
        fb1, fb2, fb3 = st.columns(3)
        a2_concrete     = fb1.checkbox("Concrete or Masonry")
        a2_cladding     = fb2.checkbox("Cladding")
        a2_curtain      = fb3.checkbox("Curtain Wall")
        a2_other_type   = st.text_input("Other façade type:", key="a2_other")

        sec("I – Works Description")
        a2_desc1 = st.text_area("I.1. Describe the façade concerned (type, materials, layers, location):", height=90, key="a2d1")
        a2_desc2 = st.text_area("I.2. Identify the waterproofing parts of the enclosure (junctions, location, material, manufacturer):", height=90, key="a2d2")
        a2_desc3 = st.text_area("I.3. Describe the junctions of the WP system with other elements:", height=80, key="a2d3")
        a2_innovative = st.radio("I.4. Innovative technique, procedure, or materials?", ["YES","NO"], index=1, horizontal=True, key="a2inn")
        if a2_innovative == "YES":
            a2_innovative_desc = st.text_area("Describe (including RD2 reference):", height=80, key="a2inndesc")
        else:
            a2_innovative_desc = ""

        sec("II – Materials of the Waterproofing System")
        da1, da2 = st.columns(2)
        a2_do_avail     = da1.radio("Delivery orders and/or certificates available?", ["YES","NO"], horizontal=True, key="a2do")
        a2_do_compliant = da2.radio("WP materials compliant with design?", ["YES","NO"], horizontal=True, key="a2comp")
        a2_do_docs      = st.text_area("Reviewed documents (delivery orders):", height=70, placeholder="1.\n2.\n3.", key="a2dodocs")
        a2_ponding      = st.radio("Water ponding test report available?", ["YES","NO"], horizontal=True, key="a2pond")
        a2_test_docs    = st.text_area("Reviewed documents (tests):", height=70, key="a2testdocs")
        a2_other_docs   = st.text_area("Other reviewed documents:", height=70, key="a2otherdocs")

        sec("III – Modifications")
        ea1, ea2 = st.columns(2)
        a2_mod          = ea1.radio("Any modification of WP system after RD0?", ["YES","NO"], index=1, horizontal=True, key="a2mod")
        a2_mod_valid    = ea2.radio("If YES — was modification validated?", ["YES","NO","N/A"], index=2, horizontal=True, key="a2modval")
        ef1, ef2 = st.columns(2)
        a2_tis_accept   = ef1.radio("Does TIS find modification acceptable?", ["YES","NO","N/A"], index=2, horizontal=True, key="a2tis")
        a2_exec_ok      = ef2.radio("Is the execution satisfactory?", ["YES","NO"], horizontal=True, key="a2exec")
        a2_reserv_refs  = st.text_input("Technical Reservation reference(s) if NO:", key="a2refs")

        sec("IV – Defects Observed")
        a2_defects      = st.radio("Defects or sign of damage regarding water infiltration on façade?", ["YES","NO"], horizontal=True, key="a2def")
        a2_defect_ref   = st.text_input("If YES — Technical Reservation ref:", key="a2defref")
        a2_defect_desc  = st.text_area("Brief description if YES:", height=80, key="a2defdesc")

        sec("V – Technical Reservations (Annex 2)")
        st.markdown("**Reservations Issued and CLOSED:**")
        a2_closed_exist = st.radio("Any closed?", ["YES","NO"], index=1, horizontal=True, key="a2cl")
        a2_closed_df    = st.data_editor(pd.DataFrame({"Concept":[""],"Issuing Cert (ref/date)":[""],"Closing Cert (ref/date)":[""]}),
                            num_rows="dynamic", use_container_width=True, key="a2cldf")
        st.markdown("**Reservations Issued and NOT CLOSED:**")
        a2_open_exist   = st.radio("Any open?", ["YES","NO"], index=1, horizontal=True, key="a2op")
        a2_open_df      = st.data_editor(pd.DataFrame({"Concept":[""],"Issuing Cert (ref/date)":[""],"Expected Closing / FINAL":[""]}),
                            num_rows="dynamic", use_container_width=True, key="a2opdf")

        sec("VI – Final Conclusion (Annex 2)")
        a2_conclusion   = st.radio("Is execution adapted to project including modifications?", ["YES","NO"], horizontal=True, key="a2concl")
        a2_concl_text   = st.text_area("Develop your answer:", height=150, key="a2concltxt")

# ═══════════════════════════════════════════════════════════════════════════════
# ANNEX 3 – BASEMENT
# ═══════════════════════════════════════════════════════════════════════════════
with tab_a3:
    if not wp_basements:
        st.info("Enable **Basements** in Tab 3 – WP Works to unlock this annex.")
    else:
        sec("Annex 3 – Waterproofing of Basement")

        sec("I – Works Description")
        a3_desc1 = st.text_area("I.1. Describe the basement concerned (type, materials, layers, location):", height=90, key="a3d1")
        a3_desc2 = st.text_area("I.2. Describe the waterproofing system of the basement:", height=90, key="a3d2")
        a3_desc3 = st.text_area("I.3. Describe the waterproofing of junctions (raft – vertical wall):", height=80, key="a3d3")
        a3_innovative = st.radio("I.4. Innovative technique, procedure, or materials?", ["YES","NO"], index=1, horizontal=True, key="a3inn")
        if a3_innovative == "YES":
            a3_innovative_desc = st.text_area("Describe (including RD2 reference):", height=80, key="a3inndesc")
        else:
            a3_innovative_desc = ""

        sec("II – Materials of the Waterproofing System")
        ga1, ga2 = st.columns(2)
        a3_do_avail     = ga1.radio("Delivery orders and/or certificates available?", ["YES","NO"], horizontal=True, key="a3do")
        a3_do_compliant = ga2.radio("WP materials compliant with design?", ["YES","NO"], horizontal=True, key="a3comp")
        a3_do_docs      = st.text_area("Reviewed documents (delivery orders):", height=70, placeholder="1.\n2.\n3.", key="a3dodocs")
        a3_ponding      = st.radio("Water ponding test report available?", ["YES","NO"], horizontal=True, key="a3pond")
        a3_test_docs    = st.text_area("Reviewed documents (tests):", height=70, key="a3testdocs")
        a3_other_docs   = st.text_area("Other reviewed documents:", height=70, key="a3otherdocs")

        sec("III – Modifications")
        ha1, ha2 = st.columns(2)
        a3_mod          = ha1.radio("Any modification of WP system after RD0?", ["YES","NO"], index=1, horizontal=True, key="a3mod")
        a3_mod_valid    = ha2.radio("If YES — was modification validated?", ["YES","NO","N/A"], index=2, horizontal=True, key="a3modval")
        hb1, hb2 = st.columns(2)
        a3_tis_accept   = hb1.radio("Does TIS find modification acceptable?", ["YES","NO","N/A"], index=2, horizontal=True, key="a3tis")
        a3_exec_ok      = hb2.radio("Is the execution satisfactory?", ["YES","NO"], horizontal=True, key="a3exec")
        a3_reserv_refs  = st.text_input("Technical Reservation reference(s) if NO:", key="a3refs")

        sec("IV – Defects Observed")
        a3_defects      = st.radio("Defects or sign of damage regarding water infiltration in basement?", ["YES","NO"], horizontal=True, key="a3def")
        a3_defect_ref   = st.text_input("If YES — Technical Reservation ref:", key="a3defref")
        a3_defect_desc  = st.text_area("Brief description if YES:", height=80, key="a3defdesc")

        sec("V – Technical Reservations (Annex 3)")
        st.markdown("**Reservations Issued and CLOSED:**")
        a3_closed_exist = st.radio("Any closed?", ["YES","NO"], index=1, horizontal=True, key="a3cl")
        a3_closed_df    = st.data_editor(pd.DataFrame({"Concept":[""],"Issuing Cert (ref/date)":[""],"Closing Cert (ref/date)":[""]}),
                            num_rows="dynamic", use_container_width=True, key="a3cldf")
        st.markdown("**Reservations Issued and NOT CLOSED:**")
        a3_open_exist   = st.radio("Any open?", ["YES","NO"], index=1, horizontal=True, key="a3op")
        a3_open_df      = st.data_editor(pd.DataFrame({"Concept":[""],"Issuing Cert (ref/date)":[""],"Expected Closing / FINAL":[""]}),
                            num_rows="dynamic", use_container_width=True, key="a3opdf")

        sec("VI – Final Conclusion (Annex 3)")
        a3_conclusion   = st.radio("Is execution adapted to project including modifications?", ["YES","NO"], horizontal=True, key="a3concl")
        a3_concl_text   = st.text_area("Develop your answer:", height=150, key="a3concltxt")

# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENT GENERATION HELPERS
# ═══════════════════════════════════════════════════════════════════════════════

BLUE        = RGBColor(0x00, 0x72, 0xBB)
DARK_NAVY   = RGBColor(0x0D, 0x2B, 0x45)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BLUE  = RGBColor(0xE8, 0xF4, 0xFD)
LIGHT_GREY  = RGBColor(0xF2, 0xF2, 0xF2)
BLACK       = RGBColor(0x00, 0x00, 0x00)

def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color="CCCCCC", size=4):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ["top","left","bottom","right"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)

def add_run(para, text, bold=False, italic=False, size=9, color=None):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run

def header_cell(cell, text, bg="0D2B45", text_size=9):
    set_cell_bg(cell, bg)
    set_cell_borders(cell, "AAAAAA")
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(3)
    para.paragraph_format.space_after  = Pt(3)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(text_size)
    run.font.color.rgb = WHITE
    return cell

def data_cell(cell, text, bold=False, size=9, bg=None, italic=False):
    if bg:
        set_cell_bg(cell, bg)
    set_cell_borders(cell, "CCCCCC")
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(2)
    run = para.add_run(str(text) if text else "")
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return cell

def label_value_cell(cell, label, value, size=9):
    set_cell_borders(cell, "CCCCCC")
    set_cell_bg(cell, "FFFFFF")
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(2)
    add_run(para, label + " ", bold=True, size=size)
    add_run(para, value or "—", size=size)

def add_section_header(doc, title, bg="0D2B45"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.columns[0].width = Inches(6.5)
    cell = tbl.cell(0, 0)
    header_cell(cell, title, bg=bg, text_size=9)
    doc.add_paragraph()

def checkbox(checked: bool) -> str:
    return "☑" if checked else "☐"

def yn_text(val, check_val="YES") -> str:
    return f"{checkbox(val == check_val)} YES    {checkbox(val != check_val)} NO"

def yn3_text(val) -> str:
    return f"{checkbox(val=='YES')} YES    {checkbox(val=='NO')} NO    {checkbox(val=='N/A')} N/A"

def add_reservations_table(doc, df: pd.DataFrame, col3_label="Expected Closing / FINAL"):
    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [Inches(2.6), Inches(2.0), Inches(1.9)]
    for i, w in enumerate(widths):
        tbl.columns[i].width = w
    headers = ["Concept", "Issuing Certificate (reference/date)", col3_label]
    for i, h in enumerate(headers):
        header_cell(tbl.cell(0, i), h, bg="0D2B45")
    for _, row in df.iterrows():
        vals = list(row)
        if all((str(v).strip() == "" or str(v) == "nan") for v in vals):
            continue
        r = tbl.add_row()
        for i, v in enumerate(vals):
            data_cell(r.cells[i], "" if str(v) == "nan" else str(v))
    if len(tbl.rows) == 1:
        r = tbl.add_row()
        for i in range(3):
            data_cell(r.cells[i], "")
    doc.add_paragraph()

def add_annex_materials(doc, do_avail, do_compliant, do_docs, ponding, test_docs, other_docs):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.columns[0].width = Inches(6.5)
    header_cell(tbl.cell(0,0), "II – MATERIALS OF THE WATERPROOFING SYSTEM")

    mat_tbl = doc.add_table(rows=4, cols=2)
    mat_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    w2 = [Inches(3.25), Inches(3.25)]
    for i, w in enumerate(w2): mat_tbl.columns[i].width = w

    label_value_cell(mat_tbl.cell(0,0), "Delivery orders / certificates available?", yn_text(do_avail))
    label_value_cell(mat_tbl.cell(0,1), "WP materials compliant with design?",       yn_text(do_compliant))

    c = mat_tbl.cell(1,0).merge(mat_tbl.cell(1,1))
    set_cell_borders(c, "CCCCCC")
    set_cell_bg(c, "FFFFFF")
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    add_run(p, "Reviewed documents (delivery orders / certificates): ", bold=True, size=9)
    add_run(p, do_docs or "—", size=9)

    label_value_cell(mat_tbl.cell(2,0), "Water ponding test report available?", yn_text(ponding))
    c2 = mat_tbl.cell(2,1); set_cell_borders(c2,"CCCCCC"); set_cell_bg(c2,"FFFFFF")

    c3 = mat_tbl.cell(3,0).merge(mat_tbl.cell(3,1))
    set_cell_borders(c3, "CCCCCC"); set_cell_bg(c3,"FFFFFF")
    p3 = c3.paragraphs[0]
    p3.paragraph_format.space_before = Pt(2); p3.paragraph_format.space_after = Pt(2)
    add_run(p3, "Reviewed documents (tests): ", bold=True, size=9)
    add_run(p3, test_docs or "—", size=9)
    if other_docs and other_docs.strip():
        p4 = c3.add_paragraph()
        p4.paragraph_format.space_before = Pt(2); p4.paragraph_format.space_after = Pt(2)
        add_run(p4, "Other reviewed documents: ", bold=True, size=9)
        add_run(p4, other_docs, size=9)
    doc.add_paragraph()

def add_annex_modifications(doc, mod, mod_valid, tis_accept, exec_ok, refs):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.columns[0].width = Inches(6.5)
    header_cell(tbl.cell(0,0), "III – MODIFICATIONS REGARDING RD0")

    mod_tbl = doc.add_table(rows=2, cols=2)
    mod_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i in range(2): mod_tbl.columns[i].width = Inches(3.25)

    label_value_cell(mod_tbl.cell(0,0), "Any modification after RD0?",      yn_text(mod))
    label_value_cell(mod_tbl.cell(0,1), "If YES — modification validated?",  yn3_text(mod_valid))
    label_value_cell(mod_tbl.cell(1,0), "TIS finds modification acceptable?",yn3_text(tis_accept))
    label_value_cell(mod_tbl.cell(1,1), "Execution satisfactory?",            yn_text(exec_ok))

    if refs and refs.strip():
        c = mod_tbl.add_row().cells[0]
        c2 = c._tc
        # merge across 2 cols done via colspan
        merged = mod_tbl.cell(len(mod_tbl.rows)-1, 0).merge(mod_tbl.cell(len(mod_tbl.rows)-1, 1))
        label_value_cell(merged, "Technical Reservation ref(s):", refs)
    doc.add_paragraph()

def add_annex_reservations(doc, closed_exist, closed_df, open_exist, open_df):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.columns[0].width = Inches(6.5)
    header_cell(tbl.cell(0,0), "V – LIST OF TECHNICAL RESERVATIONS ISSUED")

    p = doc.add_paragraph()
    add_run(p, f"Technical reservations issued and CLOSED:  {yn_text(closed_exist)}", bold=True, size=9)
    add_reservations_table(doc, closed_df, col3_label="Closing Certificate (ref/date)")

    p2 = doc.add_paragraph()
    add_run(p2, f"Technical reservations issued and NOT CLOSED:  {yn_text(open_exist)}", bold=True, size=9)
    add_reservations_table(doc, open_df, col3_label="Expected Closing / FINAL")

def add_annex_conclusion(doc, conclusion, concl_text, section_num="VI"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.columns[0].width = Inches(6.5)
    header_cell(tbl.cell(0,0), f"{section_num} – FINAL CONCLUSION (DESIGN + EXECUTION)")

    c_tbl = doc.add_table(rows=2, cols=1)
    c_tbl.columns[0].width = Inches(6.5)
    label_value_cell(c_tbl.cell(0,0), "Is the execution of works adapted to the project? Including eventual modifications?", yn_text(conclusion))
    c = c_tbl.cell(1,0)
    set_cell_borders(c,"CCCCCC"); set_cell_bg(c,"FFFFFF")
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(4)
    add_run(p, concl_text or "—", size=9)
    doc.add_paragraph()

def add_signature_block(doc, city, issue_date, expert_name):
    doc.add_paragraph()
    sig_tbl = doc.add_table(rows=3, cols=1)
    sig_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    sig_tbl.columns[0].width = Inches(6.5)

    c0 = sig_tbl.cell(0,0)
    set_cell_bg(c0,"FFFFFF"); set_cell_borders(c0,"FFFFFF")
    p0 = c0.paragraphs[0]
    add_run(p0, f"Made in {city or '___'}, {issue_date.strftime('%d %B %Y') if issue_date else '___'}", size=9, italic=True)

    for i, label in [(1, expert_name or "___"), (2, "(Signature + Ink Pad)")]:
        c = sig_tbl.cell(i,0)
        set_cell_bg(c,"FFFFFF"); set_cell_borders(c,"FFFFFF")
        p = c.paragraphs[0]; p.paragraph_format.space_before = Pt(4)
        add_run(p, label, bold=(i==1), size=9)

def add_page_break(doc):
    doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# MAIN DOCUMENT BUILDER
# ═══════════════════════════════════════════════════════════════════════════════

def build_rd3_document(data: dict) -> BytesIO:
    doc = Document()

    # ── Page setup
    section = doc.sections[0]
    section.page_width   = Inches(8.5)
    section.page_height  = Inches(11)
    section.left_margin  = section.right_margin  = Inches(1.0)
    section.top_margin   = section.bottom_margin = Inches(0.9)

    # ── Default style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(9)

    d = data

    # ────────────── HEADER BAR ──────────────
    hdr_tbl = doc.add_table(rows=1, cols=2)
    hdr_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_tbl.columns[0].width = Inches(3.5)
    hdr_tbl.columns[1].width = Inches(3.0)

    c_left = hdr_tbl.cell(0,0)
    set_cell_bg(c_left, "0D2B45")
    set_cell_borders(c_left, "0D2B45")
    p_l = c_left.paragraphs[0]
    p_l.paragraph_format.space_before = Pt(6); p_l.paragraph_format.space_after = Pt(6)
    add_run(p_l, "شركة سوكوتيك أرابيا - المملكة العربية السعودية\n", bold=True, size=8, color=WHITE)
    add_run(p_l, "SOCOTEC Arabia – KSA", bold=True, size=9, color=WHITE)

    c_right = hdr_tbl.cell(0,1)
    set_cell_bg(c_right, "FFFFFF")
    set_cell_borders(c_right, "CCCCCC")
    p_r = c_right.paragraphs[0]
    p_r.paragraph_format.space_before = Pt(4); p_r.paragraph_format.space_after = Pt(4)
    add_run(p_r, f"Ref.: {d.get('ref_num','')}\nTawuniya Visit ID: {d.get('tawuniya_id','')}", size=8)

    doc.add_paragraph()

    # ────────────── REPORT TITLE ──────────────
    title_tbl = doc.add_table(rows=1, cols=1)
    title_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    title_tbl.columns[0].width = Inches(6.5)
    c_title = title_tbl.cell(0,0)
    set_cell_bg(c_title, "0072BB")
    set_cell_borders(c_title, "0072BB")
    p_title = c_title.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(8); p_title.paragraph_format.space_after = Pt(8)
    add_run(p_title, "WATERPROOFNESS FINAL REPORT  –  RD3", bold=True, size=13, color=WHITE)
    doc.add_paragraph()

    # ────────────── DOC REFERENCE TABLE ──────────────
    ref_tbl = doc.add_table(rows=3, cols=2)
    ref_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i in range(2): ref_tbl.columns[i].width = Inches(3.25)

    header_cell(ref_tbl.cell(0,0), "DOCUMENT REFERENCE")
    header_cell(ref_tbl.cell(0,1), "TIS AGENCY")
    data_cell(ref_tbl.cell(1,0), d.get("doc_ref",""))
    data_cell(ref_tbl.cell(1,1), d.get("tis_agency","SOCOTEC Arabia"))
    header_cell(ref_tbl.cell(2,0), "VERSION")
    header_cell(ref_tbl.cell(2,1), "DATE OF ISSUE")

    ref_tbl2 = doc.add_table(rows=1, cols=2)
    ref_tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i in range(2): ref_tbl2.columns[i].width = Inches(3.25)
    data_cell(ref_tbl2.cell(0,0), d.get("version","1.0"))
    data_cell(ref_tbl2.cell(0,1), d.get("issue_date",""))

    doc.add_paragraph()

    # ────────────── RESPONSIBLE EXPERTS ──────────────
    exp_tbl = doc.add_table(rows=3, cols=4)
    exp_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    col_w = [Inches(2.0), Inches(1.5), Inches(1.5), Inches(1.5)]
    for i, w in enumerate(col_w): exp_tbl.columns[i].width = w
    for j, h in enumerate(["SURNAME, NAME","PHASE","DEGREE","SPECIALITY"]):
        header_cell(exp_tbl.cell(0,j), h)
    for j, k in enumerate(["expert_name","expert_phase","expert_degree","expert_spec"]):
        data_cell(exp_tbl.cell(1,j), d.get(k,""))
    # Author row (merged)
    merged = exp_tbl.cell(2,0).merge(exp_tbl.cell(2,3))
    set_cell_borders(merged,"CCCCCC"); set_cell_bg(merged,"F8FBFF")
    p_auth = merged.paragraphs[0]
    p_auth.paragraph_format.space_before = Pt(3); p_auth.paragraph_format.space_after = Pt(3)
    add_run(p_auth, "AUTHOR OF THIS REPORT: ", bold=True, size=8, color=DARK_NAVY)
    add_run(p_auth, f"{d.get('author_name','')}    ", size=8)
    add_run(p_auth, "PHONE: ", bold=True, size=8, color=DARK_NAVY)
    add_run(p_auth, f"{d.get('author_phone','')}    ", size=8)
    add_run(p_auth, "EMAIL: ", bold=True, size=8, color=DARK_NAVY)
    add_run(p_auth, d.get("author_email",""), size=8)

    doc.add_paragraph()

    # ────────────── SECTION I – PROJECT INFO ──────────────
    add_section_header(doc, "I – PROJECT INFORMATION")
    proj_tbl = doc.add_table(rows=5, cols=2)
    proj_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i in range(2): proj_tbl.columns[i].width = Inches(3.25)

    # Row 0: title (merged)
    merged_title = proj_tbl.cell(0,0).merge(proj_tbl.cell(0,1))
    label_value_cell(merged_title, "PROJECT TITLE / NAME:", d.get("proj_title",""))
    # Row 1: address (merged)
    merged_addr = proj_tbl.cell(1,0).merge(proj_tbl.cell(1,1))
    label_value_cell(merged_addr, "ADDRESS OF THE PREMISES:", d.get("proj_address",""))
    # Row 2
    label_value_cell(proj_tbl.cell(2,0), "REFERENCE RD0:", d.get("ref_rd0",""))
    label_value_cell(proj_tbl.cell(2,1), "PRINCIPAL / OWNER:", d.get("owner",""))
    # Row 3-4: buildings (merged)
    merged_bld = proj_tbl.cell(3,0).merge(proj_tbl.cell(4,1))
    label_value_cell(merged_bld, "BUILDINGS INCLUDED IN PROJECT AND USE:", d.get("buildings_use",""))
    doc.add_paragraph()

    # ────────────── SECTION II – WP WORKS ──────────────
    add_section_header(doc, "II – WATERPROOFING WORKS CONCERNED")
    wp_tbl = doc.add_table(rows=1, cols=3)
    wp_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i in range(3): wp_tbl.columns[i].width = Inches(6.5/3)
    for i, (key, label) in enumerate([("wp_roofs","ROOFS – Annex 1"),("wp_facades","FAÇADES – Annex 2"),("wp_basements","BASEMENTS – Annex 3")]):
        c = wp_tbl.cell(0,i); set_cell_borders(c,"CCCCCC"); set_cell_bg(c,"FFFFFF")
        p = c.paragraphs[0]; p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
        add_run(p, f"{checkbox(d.get(key, False))}  {label}", bold=True, size=9)
    doc.add_paragraph()

    # ────────────── SECTION III – DATES ──────────────
    add_section_header(doc, "III – DATES CONTROL")
    dt_tbl = doc.add_table(rows=1, cols=2)
    dt_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    dt_tbl.columns[0].width = Inches(2.5)
    dt_tbl.columns[1].width = Inches(4.0)
    header_cell(dt_tbl.cell(0,0), "III.1. OCCUPANCY")
    c_occ = dt_tbl.cell(0,1); set_cell_borders(c_occ,"CCCCCC"); set_cell_bg(c_occ,"FFFFFF")
    p_occ = c_occ.paragraphs[0]
    p_occ.paragraph_format.space_before = Pt(3); p_occ.paragraph_format.space_after = Pt(3)
    add_run(p_occ, "Date of Occupancy Certificate: ", bold=True, size=9)
    add_run(p_occ, f"{d.get('occ_date_str','')}    ", size=9)
    add_run(p_occ, f"{checkbox(d.get('occ_status','')=='Expected')} Expected    {checkbox(d.get('occ_status','')=='Confirmed')} Confirmed", size=9)

    # Visits table
    header_cell(dt_tbl.add_row().cells[0], "III.2. WP CONTROLS")
    c_vis = dt_tbl.cell(1,1); set_cell_borders(c_vis,"CCCCCC"); set_cell_bg(c_vis,"FFFFFF")
    # Build sub-table inside the cell
    visits_data = d.get("visits_df", pd.DataFrame())
    vis_inner = doc.add_table(rows=1, cols=4)
    vis_inner.alignment = WD_TABLE_ALIGNMENT.CENTER
    vw = [Inches(1.2), Inches(1.0), Inches(2.0), Inches(2.3)]
    for i, w in enumerate(vw): vis_inner.columns[i].width = w
    for i, h in enumerate(["REFERENCE","DATE","INSPECTOR","PART INSPECTED"]):
        header_cell(vis_inner.cell(0,i), h, bg="184F7A")
    for _, row in visits_data.iterrows():
        vals = list(row)
        if all((str(v).strip()=="" or str(v)=="nan") for v in vals): continue
        r = vis_inner.add_row()
        for i, v in enumerate(vals):
            data_cell(r.cells[i], "" if str(v)=="nan" else str(v))
    if len(vis_inner.rows) == 1:
        r = vis_inner.add_row()
        for i in range(4): data_cell(r.cells[i], "")
    doc.add_paragraph()

    # Defects note
    if d.get("defects_note","").strip():
        def_tbl = doc.add_table(rows=1, cols=1)
        def_tbl.columns[0].width = Inches(6.5)
        header_cell(def_tbl.cell(0,0), "DEFECTS / DISORDERS NOT REPAIRED AT DATE OF THIS REPORT", bg="184F7A")
        def_tbl2 = doc.add_table(rows=1, cols=1)
        def_tbl2.columns[0].width = Inches(6.5)
        c = def_tbl2.cell(0,0); set_cell_borders(c,"CCCCCC"); set_cell_bg(c,"FFFFFF")
        p = c.paragraphs[0]; p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
        add_run(p, d.get("defects_note",""), size=9)
        doc.add_paragraph()

    # ────────────── SECTION IV – RESERVATIONS ──────────────
    add_section_header(doc, "IV – LIST OF TECHNICAL RESERVATIONS NOT CLOSED")
    rv_tbl = doc.add_table(rows=1, cols=1)
    rv_tbl.columns[0].width = Inches(6.5)
    c_rv = rv_tbl.cell(0,0); set_cell_borders(c_rv,"CCCCCC"); set_cell_bg(c_rv,"FFFFFF")
    p_rv = c_rv.paragraphs[0]
    p_rv.paragraph_format.space_before = Pt(3); p_rv.paragraph_format.space_after = Pt(3)
    add_run(p_rv, f"Technical reservations issued and not closed:  {yn_text(d.get('reserv_exist','NO'))}", bold=True, size=9)
    add_reservations_table(doc, d.get("reserv_df", pd.DataFrame({"Concept":[""],"Issuing Cert (ref/date)":[""],"Expected Closing / FINAL":[""]})))

    # ────────────── SECTION V – FINAL CONCLUSION ──────────────
    add_annex_conclusion(doc, d.get("main_conclusion","YES"), d.get("main_concl_text",""), section_num="V")

    # ────────────── SIGNATURE BLOCK ──────────────
    add_signature_block(doc, d.get("city","Riyadh"), d.get("issue_date_obj", date.today()), d.get("expert_name",""))

    # ═══════════════ ANNEXES ═══════════════

    # ── ANNEX 1 – ROOFS ──────────────────────────────────────────────────────
    if d.get("wp_roofs"):
        add_page_break(doc)

        ann1_title = doc.add_table(rows=1, cols=1)
        ann1_title.columns[0].width = Inches(6.5)
        header_cell(ann1_title.cell(0,0), "ANNEX 1 – WATERPROOFING OF ROOFS", bg="0072BB", text_size=11)
        doc.add_paragraph()

        # Type of roofs
        roof_types = []
        if d.get("a1_roof"):           roof_types.append("Roof")
        if d.get("a1_terrace_top"):    roof_types.append("Rooftop Terrace")
        if d.get("a1_terrace_int"):    roof_types.append("Intermediate Terrace")
        if d.get("a1_patios"):         roof_types.append("Patios")
        if d.get("a1_other_type",""):  roof_types.append(d["a1_other_type"])

        type_tbl = doc.add_table(rows=1, cols=5)
        type_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        type_labels = ["Roof","Rooftop Terrace","Intermediate Terrace","Patios","Other"]
        for i, lb in enumerate(type_labels):
            type_tbl.columns[i].width = Inches(1.3)
            c = type_tbl.cell(0,i); set_cell_borders(c,"CCCCCC"); set_cell_bg(c,"FFFFFF")
            p = c.paragraphs[0]; p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
            add_run(p, f"{checkbox(lb in roof_types or (lb=='Other' and bool(d.get('a1_other_type',''))))}  {lb}", size=9)
        doc.add_paragraph()

        # Works description
        add_section_header(doc, "I – WORKS DESCRIPTION", bg="184F7A")
        for n, (key, label) in enumerate([
            ("a1_desc1","I.1. Describe the roof concerned (type, materials, layers, slope, location):"),
            ("a1_desc2","I.2. Describe WP system layers (material, manufacturer, thickness) from inside outwards:"),
            ("a1_desc3","I.3. Describe junctions of WP system with other elements:"),
        ]):
            desc_tbl = doc.add_table(rows=2, cols=1)
            desc_tbl.columns[0].width = Inches(6.5)
            header_cell(desc_tbl.cell(0,0), label, bg="2A6090", text_size=8)
            c = desc_tbl.cell(1,0); set_cell_borders(c,"CCCCCC"); set_cell_bg(c,"FFFFFF")
            p = c.paragraphs[0]; p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
            add_run(p, d.get(key,"") or "—", size=9)
        inn_tbl = doc.add_table(rows=1, cols=2)
        inn_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        inn_tbl.columns[0].width = Inches(4.0); inn_tbl.columns[1].width = Inches(2.5)
        label_value_cell(inn_tbl.cell(0,0), "I.4. Innovative technique / procedure / materials?", yn_text(d.get("a1_innovative","NO")))
        c_inn = inn_tbl.cell(0,1); set_cell_borders(c_inn,"CCCCCC"); set_cell_bg(c_inn,"FFFFFF")
        if d.get("a1_innovative","NO") == "YES":
            p_inn = c_inn.paragraphs[0]
            add_run(p_inn, d.get("a1_innovative_desc",""), size=9)
        doc.add_paragraph()

        add_annex_materials(doc, d.get("a1_do_avail","NO"), d.get("a1_do_compliant","NO"),
                            d.get("a1_do_docs",""), d.get("a1_ponding","NO"),
                            d.get("a1_test_docs",""), d.get("a1_other_docs",""))
        add_annex_modifications(doc, d.get("a1_mod","NO"), d.get("a1_mod_valid","N/A"),
                                d.get("a1_tis_accept","N/A"), d.get("a1_exec_ok","YES"),
                                d.get("a1_reserv_refs",""))

        # IV – Roof WP tests
        add_section_header(doc, "IV – ROOF WATERPROOFING TESTS", bg="184F7A")
        test_tbl = doc.add_table(rows=1, cols=2)
        test_tbl.columns[0].width = Inches(4.5); test_tbl.columns[1].width = Inches(2.0)
        label_value_cell(test_tbl.cell(0,0), "Defects or damage regarding water roof infiltration?", yn_text(d.get("a1_defects","NO")))
        label_value_cell(test_tbl.cell(0,1), "If YES – TR Ref:", d.get("a1_defect_ref",""))
        doc.add_paragraph()

        add_annex_reservations(doc, d.get("a1_closed_exist","NO"), d.get("a1_closed_df", pd.DataFrame()),
                               d.get("a1_open_exist","NO"), d.get("a1_open_df", pd.DataFrame()))
        add_annex_conclusion(doc, d.get("a1_conclusion","YES"), d.get("a1_concl_text",""))
        add_signature_block(doc, d.get("city",""), d.get("issue_date_obj", date.today()), d.get("expert_name",""))

    # ── ANNEX 2 – FAÇADE ─────────────────────────────────────────────────────
    if d.get("wp_facades"):
        add_page_break(doc)
        ann2_title = doc.add_table(rows=1, cols=1)
        ann2_title.columns[0].width = Inches(6.5)
        header_cell(ann2_title.cell(0,0), "ANNEX 2 – WATERPROOFING OF FAÇADE", bg="0072BB", text_size=11)
        doc.add_paragraph()

        facade_types = []
        if d.get("a2_concrete"): facade_types.append("Concrete or Masonry")
        if d.get("a2_cladding"): facade_types.append("Cladding")
        if d.get("a2_curtain"):  facade_types.append("Curtain Wall")
        if d.get("a2_other_type",""): facade_types.append(d["a2_other_type"])

        ft_tbl = doc.add_table(rows=1, cols=4)
        ft_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, lb in enumerate(["Concrete / Masonry","Cladding","Curtain Wall","Other"]):
            ft_tbl.columns[i].width = Inches(1.625)
            c = ft_tbl.cell(0,i); set_cell_borders(c,"CCCCCC"); set_cell_bg(c,"FFFFFF")
            p = c.paragraphs[0]; p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
            chk = lb in facade_types or (lb=="Other" and bool(d.get("a2_other_type","")))
            add_run(p, f"{checkbox(chk)}  {lb}", size=9)
        doc.add_paragraph()

        add_section_header(doc, "I – WORKS DESCRIPTION", bg="184F7A")
        for key, label in [
            ("a2_desc1","I.1. Describe the façade (type, materials, layers, location):"),
            ("a2_desc2","I.2. Identify waterproofing parts of the enclosure (junctions, location, material, manufacturer):"),
            ("a2_desc3","I.3. Describe junctions of WP system with other elements:"),
        ]:
            d2_tbl = doc.add_table(rows=2, cols=1)
            d2_tbl.columns[0].width = Inches(6.5)
            header_cell(d2_tbl.cell(0,0), label, bg="2A6090", text_size=8)
            c = d2_tbl.cell(1,0); set_cell_borders(c,"CCCCCC"); set_cell_bg(c,"FFFFFF")
            p = c.paragraphs[0]; p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
            add_run(p, d.get(key,"") or "—", size=9)

        inn2_tbl = doc.add_table(rows=1, cols=2)
        inn2_tbl.columns[0].width = Inches(4.0); inn2_tbl.columns[1].width = Inches(2.5)
        label_value_cell(inn2_tbl.cell(0,0), "I.4. Innovative technique / procedure / materials?", yn_text(d.get("a2_innovative","NO")))
        c_inn2 = inn2_tbl.cell(0,1); set_cell_borders(c_inn2,"CCCCCC"); set_cell_bg(c_inn2,"FFFFFF")
        if d.get("a2_innovative","NO") == "YES":
            p_inn2 = c_inn2.paragraphs[0]; add_run(p_inn2, d.get("a2_innovative_desc",""), size=9)
        doc.add_paragraph()

        add_annex_materials(doc, d.get("a2_do_avail","NO"), d.get("a2_do_compliant","NO"),
                            d.get("a2_do_docs",""), d.get("a2_ponding","NO"),
                            d.get("a2_test_docs",""), d.get("a2_other_docs",""))
        add_annex_modifications(doc, d.get("a2_mod","NO"), d.get("a2_mod_valid","N/A"),
                                d.get("a2_tis_accept","N/A"), d.get("a2_exec_ok","YES"),
                                d.get("a2_reserv_refs",""))

        add_section_header(doc, "IV – DEFECTS OBSERVED", bg="184F7A")
        def2_tbl = doc.add_table(rows=2, cols=2)
        def2_tbl.columns[0].width = Inches(4.0); def2_tbl.columns[1].width = Inches(2.5)
        label_value_cell(def2_tbl.cell(0,0), "Defects / damage regarding water infiltration on façade?", yn_text(d.get("a2_defects","NO")))
        label_value_cell(def2_tbl.cell(0,1), "If YES – TR Ref:", d.get("a2_defect_ref",""))
        merged_desc2 = def2_tbl.cell(1,0).merge(def2_tbl.cell(1,1))
        label_value_cell(merged_desc2, "Brief description:", d.get("a2_defect_desc",""))
        doc.add_paragraph()

        add_annex_reservations(doc, d.get("a2_closed_exist","NO"), d.get("a2_closed_df", pd.DataFrame()),
                               d.get("a2_open_exist","NO"), d.get("a2_open_df", pd.DataFrame()))
        add_annex_conclusion(doc, d.get("a2_conclusion","YES"), d.get("a2_concl_text",""))
        add_signature_block(doc, d.get("city",""), d.get("issue_date_obj", date.today()), d.get("expert_name",""))

    # ── ANNEX 3 – BASEMENT ───────────────────────────────────────────────────
    if d.get("wp_basements"):
        add_page_break(doc)
        ann3_title = doc.add_table(rows=1, cols=1)
        ann3_title.columns[0].width = Inches(6.5)
        header_cell(ann3_title.cell(0,0), "ANNEX 3 – WATERPROOFING OF BASEMENT", bg="0072BB", text_size=11)
        doc.add_paragraph()

        add_section_header(doc, "I – WORKS DESCRIPTION", bg="184F7A")
        for key, label in [
            ("a3_desc1","I.1. Describe the basement concerned (type, materials, layers, location):"),
            ("a3_desc2","I.2. Describe the waterproofing system of the basement:"),
            ("a3_desc3","I.3. Describe the waterproofing of junctions (raft – vertical wall):"),
        ]:
            d3_tbl = doc.add_table(rows=2, cols=1)
            d3_tbl.columns[0].width = Inches(6.5)
            header_cell(d3_tbl.cell(0,0), label, bg="2A6090", text_size=8)
            c = d3_tbl.cell(1,0); set_cell_borders(c,"CCCCCC"); set_cell_bg(c,"FFFFFF")
            p = c.paragraphs[0]; p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
            add_run(p, d.get(key,"") or "—", size=9)

        inn3_tbl = doc.add_table(rows=1, cols=2)
        inn3_tbl.columns[0].width = Inches(4.0); inn3_tbl.columns[1].width = Inches(2.5)
        label_value_cell(inn3_tbl.cell(0,0), "I.4. Innovative technique / procedure / materials?", yn_text(d.get("a3_innovative","NO")))
        c_inn3 = inn3_tbl.cell(0,1); set_cell_borders(c_inn3,"CCCCCC"); set_cell_bg(c_inn3,"FFFFFF")
        if d.get("a3_innovative","NO") == "YES":
            p_inn3 = c_inn3.paragraphs[0]; add_run(p_inn3, d.get("a3_innovative_desc",""), size=9)
        doc.add_paragraph()

        add_annex_materials(doc, d.get("a3_do_avail","NO"), d.get("a3_do_compliant","NO"),
                            d.get("a3_do_docs",""), d.get("a3_ponding","NO"),
                            d.get("a3_test_docs",""), d.get("a3_other_docs",""))
        add_annex_modifications(doc, d.get("a3_mod","NO"), d.get("a3_mod_valid","N/A"),
                                d.get("a3_tis_accept","N/A"), d.get("a3_exec_ok","YES"),
                                d.get("a3_reserv_refs",""))

        add_section_header(doc, "IV – DEFECTS OBSERVED", bg="184F7A")
        def3_tbl = doc.add_table(rows=2, cols=2)
        def3_tbl.columns[0].width = Inches(4.0); def3_tbl.columns[1].width = Inches(2.5)
        label_value_cell(def3_tbl.cell(0,0), "Defects / damage regarding water infiltration in basement?", yn_text(d.get("a3_defects","NO")))
        label_value_cell(def3_tbl.cell(0,1), "If YES – TR Ref:", d.get("a3_defect_ref",""))
        merged_desc3 = def3_tbl.cell(1,0).merge(def3_tbl.cell(1,1))
        label_value_cell(merged_desc3, "Brief description:", d.get("a3_defect_desc",""))
        doc.add_paragraph()

        add_annex_reservations(doc, d.get("a3_closed_exist","NO"), d.get("a3_closed_df", pd.DataFrame()),
                               d.get("a3_open_exist","NO"), d.get("a3_open_df", pd.DataFrame()))
        add_annex_conclusion(doc, d.get("a3_conclusion","YES"), d.get("a3_concl_text",""))
        add_signature_block(doc, d.get("city",""), d.get("issue_date_obj", date.today()), d.get("expert_name",""))

    # ── Footer on all pages ──
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.paragraph_format.space_before = Pt(2)
    add_run(footer_para, "SOCOTEC Arabia – KSA  |  6897 King Fahd Road, Al Olaya, Riyadh  |  Restricted", size=7, color=RGBColor(0x88,0x88,0x88))

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ═══════════════════════════════════════════════════════════════════════════════
# GENERATE BUTTON
# ═══════════════════════════════════════════════════════════════════════════════
st.markdown("---")
col_btn1, col_btn2, col_btn3 = st.columns([1,2,1])
with col_btn2:
    generate = st.button("⬇️  Generate RD3 Report (.docx)", use_container_width=True, type="primary")

if generate:
    # Collect all form data into one dict
    # (variables are read from session / local scope – all defined above)
    try:
        data_dict = {
            # Report Info
            "ref_num": ref_num, "tawuniya_id": tawuniya_id, "doc_ref": doc_ref,
            "tis_agency": tis_agency, "version": version,
            "issue_date": issue_date.strftime("%d/%m/%Y"), "issue_date_obj": issue_date,
            "expert_name": expert_name, "expert_phase": expert_phase,
            "expert_degree": expert_degree, "expert_spec": expert_spec,
            "author_name": author_name, "author_phone": author_phone,
            "author_email": author_email, "city": city,
            # Project
            "proj_title": proj_title, "proj_address": proj_address,
            "ref_rd0": ref_rd0, "owner": owner, "buildings_use": buildings_use,
            # WP Works
            "wp_roofs": wp_roofs, "wp_facades": wp_facades, "wp_basements": wp_basements,
            # Dates
            "occ_date_str": occ_date_str, "occ_status": occ_status,
            "visits_df": visits_df, "defects_note": defects_note,
            # Reservations
            "reserv_exist": reserv_exist, "reserv_df": reserv_df,
            # Conclusion
            "main_conclusion": main_conclusion, "main_concl_text": main_concl_text,
        }

        # Annex 1 (conditional)
        if wp_roofs:
            data_dict.update({
                "a1_roof": a1_roof, "a1_terrace_top": a1_terrace_top,
                "a1_terrace_int": a1_terrace_int, "a1_patios": a1_patios,
                "a1_other_type": a1_other_type, "a1_desc1": a1_desc1,
                "a1_desc2": a1_desc2, "a1_desc3": a1_desc3,
                "a1_innovative": a1_innovative, "a1_innovative_desc": a1_innovative_desc,
                "a1_do_avail": a1_do_avail, "a1_do_compliant": a1_do_compliant,
                "a1_do_docs": a1_do_docs, "a1_ponding": a1_ponding,
                "a1_test_docs": a1_test_docs, "a1_other_docs": a1_other_docs,
                "a1_mod": a1_mod, "a1_mod_valid": a1_mod_valid,
                "a1_tis_accept": a1_tis_accept, "a1_exec_ok": a1_exec_ok,
                "a1_reserv_refs": a1_reserv_refs, "a1_defects": a1_defects,
                "a1_defect_ref": a1_defect_ref,
                "a1_closed_exist": a1_closed_exist, "a1_closed_df": a1_closed_df,
                "a1_open_exist": a1_open_exist, "a1_open_df": a1_open_df,
                "a1_conclusion": a1_conclusion, "a1_concl_text": a1_concl_text,
            })

        # Annex 2 (conditional)
        if wp_facades:
            data_dict.update({
                "a2_concrete": a2_concrete, "a2_cladding": a2_cladding,
                "a2_curtain": a2_curtain, "a2_other_type": a2_other_type,
                "a2_desc1": a2_desc1, "a2_desc2": a2_desc2, "a2_desc3": a2_desc3,
                "a2_innovative": a2_innovative, "a2_innovative_desc": a2_innovative_desc,
                "a2_do_avail": a2_do_avail, "a2_do_compliant": a2_do_compliant,
                "a2_do_docs": a2_do_docs, "a2_ponding": a2_ponding,
                "a2_test_docs": a2_test_docs, "a2_other_docs": a2_other_docs,
                "a2_mod": a2_mod, "a2_mod_valid": a2_mod_valid,
                "a2_tis_accept": a2_tis_accept, "a2_exec_ok": a2_exec_ok,
                "a2_reserv_refs": a2_reserv_refs, "a2_defects": a2_defects,
                "a2_defect_ref": a2_defect_ref, "a2_defect_desc": a2_defect_desc,
                "a2_closed_exist": a2_closed_exist, "a2_closed_df": a2_closed_df,
                "a2_open_exist": a2_open_exist, "a2_open_df": a2_open_df,
                "a2_conclusion": a2_conclusion, "a2_concl_text": a2_concl_text,
            })

        # Annex 3 (conditional)
        if wp_basements:
            data_dict.update({
                "a3_desc1": a3_desc1, "a3_desc2": a3_desc2, "a3_desc3": a3_desc3,
                "a3_innovative": a3_innovative, "a3_innovative_desc": a3_innovative_desc,
                "a3_do_avail": a3_do_avail, "a3_do_compliant": a3_do_compliant,
                "a3_do_docs": a3_do_docs, "a3_ponding": a3_ponding,
                "a3_test_docs": a3_test_docs, "a3_other_docs": a3_other_docs,
                "a3_mod": a3_mod, "a3_mod_valid": a3_mod_valid,
                "a3_tis_accept": a3_tis_accept, "a3_exec_ok": a3_exec_ok,
                "a3_reserv_refs": a3_reserv_refs, "a3_defects": a3_defects,
                "a3_defect_ref": a3_defect_ref, "a3_defect_desc": a3_defect_desc,
                "a3_closed_exist": a3_closed_exist, "a3_closed_df": a3_closed_df,
                "a3_open_exist": a3_open_exist, "a3_open_df": a3_open_df,
                "a3_conclusion": a3_conclusion, "a3_concl_text": a3_concl_text,
            })

        with st.spinner("Generating document…"):
            buf = build_rd3_document(data_dict)

        fname = f"RD3_{proj_title.replace(' ','_') if proj_title else 'Report'}_{issue_date.strftime('%Y%m%d')}.docx"
        st.success("✅ Report generated successfully!")
        st.download_button(
            label="📥 Download RD3 Report (.docx)",
            data=buf,
            file_name=fname,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

    except Exception as e:
        st.error(f"❌ Generation failed: {e}")
        st.exception(e)

st.markdown("---")
if st.button("← Back to Home", key="back_home"):
    st.switch_page("Home.py")

