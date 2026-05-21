"""
rd3_extractor.py
Smart document extraction for RD3 Waterproofness Report generator.
Extracts data from:
  - Waterproofing warranty certificates (Arabic PDF)
  - SOCOTEC TIS visit reports (DOCX)
  - Architectural drawings (PDF)

Completely separate from rd6_extractor.py — no risk to RD6 app.
"""
import re
import io
import unicodedata as _ucd
import zipfile as _zf
import pdfplumber


# ═══════════════════════════════════════════════════════════════════════════════
# ── Smart Document Extraction (Warranty + Visit Report + Arch Drawings) ────────
# ═══════════════════════════════════════════════════════════════════════════════

import unicodedata as _ucd
import zipfile as _zf


def _fix_ar(text):
    """Fix visually-ordered Arabic from PDFs (presentation forms + RTL)."""
    if not text: return text
    norm = _ucd.normalize('NFKC', text)
    has_ar = any('\u0600' <= c <= '\u06ff' for c in norm)
    if not has_ar: return norm
    words = norm.strip().split()
    result = []
    for w in reversed(words):
        ar_ratio = sum(1 for c in w if '\u0600' <= c <= '\u06ff') / max(len(w), 1)
        result.append(''.join(reversed(w)) if ar_ratio > 0.4 else w)
    return ' '.join(result)


def extract_from_warranty_pdf(path_or_bytes):
    """
    Extract from waterproofing warranty certificate PDF.
    Returns: owner, plot_no, plan_no, wp_type_arabic, wp_type_description,
             location, warranty_start, contractor, doc_no
    """
    src = io.BytesIO(path_or_bytes) if isinstance(path_or_bytes, (bytes, bytearray)) else path_or_bytes
    with pdfplumber.open(src) as pdf:
        raw = '\n'.join(pg.extract_text() or '' for pg in pdf.pages)

    fixed = '\n'.join(_fix_ar(line) for line in raw.split('\n'))
    norm = _ucd.normalize('NFKC', raw)
    result = {}

    # Owner
    m = re.search(r'المالك\s*/\s*([^/\n]{5,50}?)(?:\s*/|\s*$)', fixed)
    if m: result['owner'] = m.group(1).strip()

    # Plot number (digits correct in norm before char reversal)
    m = re.search(r'(\d{3,})/ب', norm) or re.search(r'ب/(\d{3,})', norm)
    if m:
        n = m.group(1)
        result['plot_no'] = n + '/ب'
    else:
        m = re.search(r'رقم القطعة\s*/\s*(\S+)', fixed)
        if m: result['plot_no'] = m.group(1).strip()

    # Plan number
    m = re.search(r'رقم المخطط\s*/\s*(\S+)', fixed)
    if m: result['plan_no'] = m.group(1).strip()

    # WP type → key I.2 input
    m = re.search(r'نوع العزل\s*/\s*([^\n]+)', fixed)
    if m:
        wp = re.sub(r'\b(\d{2,4})\b', lambda x: x.group()[::-1], m.group(1).strip())
        result['wp_type_arabic'] = wp
        result['wp_type_description'] = _translate_wp_type(wp)

    # Location
    m = re.search(r'الموقع[:\s]+([^\n]+)', fixed)
    if m: result['location'] = m.group(1).strip()

    # Warranty start date (18/11/2025 format in norm)
    # Try multiple patterns, prefer earliest date
    dates = re.findall(r'\b(\d{1,2}/\d{1,2}/\d{4})\b', norm)
    if len(dates) >= 2:
        result['warranty_start'] = dates[-2]   # usually 2nd-to-last date is the warranty start
    elif dates:
        result['warranty_start'] = dates[-1]

    # Contractor
    for pat in [r'^((?:مؤسسة|شركة)[^\n/]{5,60})$', r'((?:مؤسسة|شركة)[^\n/]{5,60})']:
        m = re.search(pat, fixed, re.MULTILINE)
        if m:
            c = re.sub(r'\s+[A-Z][a-z].*', '', m.group(1)).strip()
            if len(c) > 8: result['contractor'] = c; break

    # Document number
    m = re.search(r'(\d{8,12})', norm)
    if m: result['doc_no'] = m.group(1)

    return result


def _translate_wp_type(wp_ar):
    """Convert Arabic WP system to English I.2 description."""
    parts = []
    if any(w in wp_ar for w in ['صبة رغوية', 'صبة رغویة', 'ميول', 'میول']):
        parts.append('a sloped concrete screed directed to the drainage outlet')
    layer = '2' if any(w in wp_ar for w in ['طبقتین', 'طبقتين']) else \
            '3' if 'ثلاث' in wp_ar else '1'
    rn = re.search(r'(?:رول|Roll)\s*(\d{3})', wp_ar)
    if 'رول' in wp_ar or 'غشاء' in wp_ar:
        roll = ' (Roll {})'.format(rn.group(1)) if rn else ''
        parts.append('{} layer(s) of bituminous waterproofing membrane{}'.format(layer, roll))
    if 'تسلیح' in wp_ar or 'تسليح' in wp_ar:
        rn2 = re.search(r'(?:تسلیح|تسليح)\s*(\d+)', wp_ar)
        parts.append('reinforced waterproofing roll{}'.format(' '+rn2.group(1) if rn2 else ''))
    ti = re.search(r'(\d+)\s*سم\s*حراري|حراري\s*(\d+)\s*سم', wp_ar)
    if ti or 'حراري' in wp_ar:
        cm = (ti.group(1) or ti.group(2)) if ti else '4'
        parts.append('{}cm thermal insulation'.format(cm))
    if 'رغوية' in wp_ar or 'رغویة' in wp_ar:
        if not any('sloped' in p for p in parts):
            parts.append('polyurethane foam spray layer')
    base = 'The waterproofing system consists of ' if parts else \
           'The waterproofing system consists of: ' + wp_ar + ' (from warranty certificate)'
    return (base + ', '.join(parts) + '.') if parts else base


def extract_from_visit_docx(path_or_bytes):
    """
    Extract visit data + images from SOCOTEC TIS visit report DOCX.
    Returns: malath_ref, socotec_ref, project_name, address, visit_ref,
             visit_no, visit_subject, visit_date, inspector_name,
             inspector_email, images [(fname, bytes), ...]
    """
    from docx import Document
    if isinstance(path_or_bytes, (bytes, bytearray)):
        src = io.BytesIO(path_or_bytes); raw_b = path_or_bytes
    else:
        src = path_or_bytes
        with open(path_or_bytes, 'rb') as f: raw_b = f.read()

    doc = Document(io.BytesIO(raw_b) if isinstance(raw_b, (bytes, bytearray)) else src)
    data = {'images': []}

    FM = {
        'Project Malath Ref': 'malath_ref',
        'Project Socotec Ref': 'socotec_ref',
        'Project Name': 'project_name',
        'Project Address': 'address',
        'Visit Malath ID': 'visit_malath_id',
        'Visit Socotec Ref': 'visit_ref',
        'Visit No': 'visit_no',
        'Visit Subject': 'visit_subject',
        'Visit Date': 'visit_date',
        'Inspector Name': 'inspector_name',
        'Inspector Email': 'inspector_email',
    }
    if doc.tables:
        for row in doc.tables[0].rows:
            cells = row.cells
            if len(cells) >= 2:
                lbl = cells[0].text.strip()
                val = cells[1].text.strip()
                if lbl in FM and val:
                    data[FM[lbl]] = val

    try:
        with _zf.ZipFile(io.BytesIO(raw_b)) as z:
            for name in z.namelist():
                if name.startswith('word/media/') and \
                   any(name.lower().endswith(e) for e in ['.jpg','.jpeg','.png','.gif']):
                    data['images'].append((name.split('/')[-1], z.read(name)))
    except Exception:
        pass
    return data


def extract_facade_from_arch_pdf(path_or_bytes):
    """
    Extract facade material info from architectural PDF elevation drawings.
    Returns: facade_description (English I.1), building_type, floors
    """
    src = io.BytesIO(path_or_bytes) if isinstance(path_or_bytes, (bytes, bytearray)) else path_or_bytes
    with pdfplumber.open(src) as pdf:
        raw = '\n'.join(pg.extract_text() or '' for pg in pdf.pages)
    norm = _ucd.normalize('NFKC', raw)

    MATS = [('حجر', 'stone cladding on main façade'),
            ('دهان رمادي', 'gray painted plaster'),
            ('دهان ابيض', 'white painted plaster'),
            ('زجاج', 'aluminum-framed glazed windows and openings'),
            ('الومنيوم', 'aluminum frames')]

    found = [en for ar, en in MATS if ar in norm]
    if not found:
        found = ['plastered block walls', 'aluminum glazed windows and openings']

    floors = 'G+1'
    if 'G+2' in raw or 'دورين' in norm: floors = 'G+2'

    desc = 'The facades consist of ' + ', '.join(found) + '.'
    return {
        'facade_description': desc,
        'building_type': 'commercial' if any(w in norm for w in ['تجاري','مكتبي']) else 'residential',
        'floors': floors,
    }


def lookup_malath_from_tawuniya(excel_path, taw_pol_no):
    """
    Reverse lookup: given a Tawuniya policy number, return the Malath IDI number.
    Uses the Tuw-Mlth sheet in malath_log.xlsx.
    Returns str IDI or None if not found.
    """
    import openpyxl
    try:
        target = str(taw_pol_no).strip().replace('.0', '')
        wb = openpyxl.load_workbook(excel_path, data_only=True)
        if 'Tuw-Mlth' not in wb.sheetnames:
            return None
        ws = wb['Tuw-Mlth']
        for row in ws.iter_rows(min_row=2, values_only=True):
            # row[1] = Malath IDI, row[2] = Tawuniya policy
            taw = str(row[2]).strip().replace('.0', '') if row[2] is not None else ''
            if taw == target:
                idi = str(row[1]).strip().replace('.0', '') if row[1] is not None else ''
                return idi if idi else None
    except Exception:
        pass
    return None
