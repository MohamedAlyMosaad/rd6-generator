"""
rd3_generator.py — Fills Template_RD3.docx with project data.
Matches the AHA-RD3 reference sample formatting exactly.

Signature block (body paragraphs × 4 blocks):
  (RESPONSIBLE EXPERTS NAMES)           → Author: header row (3-column, space-padded)
  (RESPONSIBLE EXPERTS SIGNATURES + INK PAD) → Eng. {name} / Reviewer: row + optional sig image
  (RESPONSIBLE EXPERTS POSITIONS)       → Eng. {reviewer} / Eng. {manager} row

SDT checkboxes (0-based position):
  0=ROOFS 1=FACADES 2=BASEMENTS
  3-7  = Annex1 roof type
  8-13 = Annex1 material Y/N
  14-17= Annex2 facade type
  18-23= Annex2 material Y/N
  24-29= Annex3 material Y/N
"""
import io
from lxml import etree
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement

W   = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
W14 = 'http://schemas.microsoft.com/office/word/2010/wordml'
XML_SPACE = '{http://www.w3.org/XML/1998/namespace}space'

# SOCOTEC dark blue — used for all filled-in values
BLUE = RGBColor(0x2F, 0x54, 0x96)

# ── Low-level XML helpers ─────────────────────────────────────────────────────

def _get_text(elem):
    return ''.join(t.text or '' for t in elem.iter(f'{{{W}}}t'))


def _set_para_text(para_elem, new_text):
    """Replace ALL text in a paragraph element with new_text (single run)."""
    t_elems = list(para_elem.iter(f'{{{W}}}t'))
    if not t_elems:
        r = etree.SubElement(para_elem, f'{{{W}}}r')
        t = etree.SubElement(r, f'{{{W}}}t')
        t.text = new_text
        if new_text and (new_text[0] == ' ' or new_text[-1] == ' '):
            t.set(XML_SPACE, 'preserve')
    else:
        t_elems[0].text = new_text
        for t in t_elems[1:]:
            t.text = ''
        if new_text and (new_text[0] == ' ' or new_text[-1] == ' '):
            t_elems[0].set(XML_SPACE, 'preserve')


def _replace_global(body, old_text, new_text):
    """Replace paragraphs containing old_text anywhere in the document."""
    for p in body.iter(f'{{{W}}}p'):
        full = _get_text(p)
        if old_text in full:
            _set_para_text(p, full.replace(old_text, new_text))


def _toggle_sdt(body, index, checked):
    """Toggle nth SDT checkbox (0-based)."""
    sdts = list(body.iter(f'{{{W}}}sdt'))
    if index >= len(sdts):
        return
    sdt = sdts[index]
    chk = sdt.find(f'.//{{{W14}}}checked')
    if chk is not None:
        chk.set(f'{{{W14}}}val', '1' if checked else '0')
    content = sdt.find(f'{{{W}}}sdtContent')
    if content is not None:
        for t in content.iter(f'{{{W}}}t'):
            t.text = '\u2612' if checked else '\u2610'


def _append_to_label_para(cell_elem, label, value):
    """Find paragraph starting with label and append value to it (handles multi-run)."""
    for p in cell_elem.findall(f'.//{{{W}}}p'):
        full_txt = _get_text(p).strip()
        if full_txt.startswith(label.strip()):
            combined = full_txt.rstrip() + (' ' + str(value) if value else '')
            _set_para_text(p, combined)
            return


def _fill_cell_text_tc(tc_elem, text):
    """Fill a table cell (w:tc) with plain text."""
    paras = tc_elem.findall(f'.//{{{W}}}p')
    if not paras:
        return
    for t in paras[0].iter(f'{{{W}}}t'):
        t.text = ''
    r = etree.SubElement(paras[0], f'{{{W}}}r')
    t = etree.SubElement(r, f'{{{W}}}t')
    t.text = str(text) if text else ''
    if text and (str(text)[0] == ' ' or str(text)[-1] == ' '):
        t.set(XML_SPACE, 'preserve')


def _append_text_to_cell_last_para(tc_elem, text):
    """Append text after 'Please expose...' paragraph in a cell."""
    paras = tc_elem.findall(f'.//{{{W}}}p')
    for i, p in enumerate(paras):
        if 'Please expose' in _get_text(p) or 'develop' in _get_text(p).lower():
            target = paras[i + 1] if i + 1 < len(paras) else paras[-1]
            r = etree.SubElement(target, f'{{{W}}}r')
            te = etree.SubElement(r, f'{{{W}}}t')
            te.text = text
            te.set(XML_SPACE, 'preserve')
            return
    if paras:
        r = etree.SubElement(paras[-1], f'{{{W}}}r')
        te = etree.SubElement(r, f'{{{W}}}t')
        te.text = text
        te.set(XML_SPACE, 'preserve')


# ── python-docx cell helper (uses high-level API for reliable fill) ───────────

def _fill_expert_cell(table, row_idx, col_idx, text):
    """Fill expert cell — removes extra empty paragraphs so all cells stay at equal height."""
    cell = table.rows[row_idx].cells[col_idx]
    tc = cell._tc
    # Remove all paragraphs except the first (extra empty paras cause misaligned rows)
    all_paras = tc.findall(f'.//{{{W}}}p')
    for p in all_paras[1:]:
        parent = p.getparent()
        if parent is not None and parent.tag != f'{{{W}}}tbl':
            parent.remove(p)
    cell.paragraphs[0].clear()
    run = cell.paragraphs[0].add_run(text)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = BLUE


def _fill_label_value_cell(table, row_idx, col_idx, label_value_dict):
    """Fill a cell with label: value pairs using python-docx API with proper styling."""
    cell = table.rows[row_idx].cells[col_idx]
    for para in cell.paragraphs:
        txt = para.text.strip()
        for label, value in label_value_dict.items():
            if txt.startswith(label.strip(':').strip()):
                para.clear()
                run_label = para.add_run(txt.rstrip() + ' ')
                run_label.bold = True
                run_label.font.size = Pt(9)
                if value:
                    run_val = para.add_run(str(value))
                    run_val.bold = True
                    run_val.font.size = Pt(9)
                    run_val.font.color.rgb = BLUE
                break


# ── Reference builder ─────────────────────────────────────────────────────────

def build_rd3_reference(eng_full, nt_ft, idi_no, taw_pol, ins_type='Malath'):
    """
    Build full RD3 reference and short ref.
    Initials: first letter of first name + first 2 letters of last name (uppercase).
    Mohamed Mossad → MMO, Abdullah Al Habib → AAL
    """
    parts = eng_full.strip().split()
    if len(parts) >= 2:
        initials = (parts[0][0] + parts[-1][:2]).upper()   # first[0] + last[:2]
    elif parts:
        initials = parts[0][:3].upper()
    else:
        initials = 'ENG'

    idi = str(idi_no).strip().replace('.0', '')
    taw = str(taw_pol).strip().replace('.0', '')
    nt  = str(nt_ft).strip()

    if ins_type == 'Tawuniya':
        full_ref  = '{}-RD3-{}-01'.format(initials, taw) if taw else '{}-RD3-{}-01'.format(initials, idi)
        short_ref = taw if taw else idi
    else:
        if taw:
            full_ref  = '{}-RD3-{}{}-{}-1'.format(initials, nt, idi, taw)
            short_ref = '{}{}-{}'.format(nt, idi, taw)
        else:
            full_ref  = '{}-RD3-{}{}-1'.format(initials, nt, idi)
            short_ref = '{}{}'.format(nt, idi)

    return full_ref, short_ref


# ── Signature block filling ───────────────────────────────────────────────────

def _make_sig_line(eng_full, reviewer, manager):
    """
    Build the 3 signature paragraph texts matching AHA layout exactly.
    Uses space-padding to create 3-column visual layout.
    Col1 (~pos 0): Author / Eng. name
    Col2 (~pos 76): ENGINEER IN CHARGE / Reviewer:
    Col3 (~pos 131): HEAD OF LOCAL DEPARTMENT / OR MANAGER / Eng. manager
    """
    # Line 1: Author header (154 chars matching AHA)
    author_line = '\tAuthor:' + ' ' * 78 + 'ENGINEER IN CHARGE' + ' ' * 26 + 'HEAD OF LOCAL DEPARTMENT'

    # Line 2: Eng. name + Reviewer label (155 chars)
    eng_str = 'Eng. ' + eng_full          # variable length
    pad1 = max(1, 76 - len(eng_str))       # space between name and Reviewer:
    pad2 = max(1, 155 - len(eng_str) - pad1 - 9 - 10)  # space before OR MANAGER
    name_line = eng_str + ' ' * pad1 + 'Reviewer:' + ' ' * pad2 + 'OR MANAGER'

    # Line 3: Reviewer name + Manager (167 chars)
    reviewer_str = 'Eng. ' + reviewer
    pad3 = 65                              # leading spaces (fixed)
    pad4 = max(1, 131 - pad3 - len(reviewer_str))  # space before manager
    manager_str = 'Eng. ' + manager.upper()
    reviewer_line = ' ' * pad3 + reviewer_str + ' ' * pad4 + manager_str + ' ' * 19

    return author_line, name_line, reviewer_line


def _fill_signature_blocks(body, data):
    """Replace all 4 signature placeholder blocks in body paragraphs."""
    eng_full  = data.get('eng_full', '')
    reviewer  = data.get('reviewer_name', '')
    manager   = data.get('manager_name', 'Nizar Lazreg')
    issue_str = data.get('issue_date', '')
    city      = data.get('city', 'Riyadh')

    # Format date dd-mm-yyyy
    try:
        parts = issue_str.split('/')
        issue_fmt = '{:02d}/{:02d}/{}'.format(int(parts[0]), int(parts[1]), parts[2])
    except Exception:
        issue_fmt = issue_str

    # Build signature lines
    author_line, name_line, reviewer_line = _make_sig_line(eng_full, reviewer, manager)

    # Replace date paragraphs (preserves any leading \t\t prefix)
    _replace_global(body, 'Made in (CITY), (DATE)',
                    'Issued in {} on {}'.format(city, issue_fmt))

    # Replace signature placeholders
    _replace_global(body, '(RESPONSIBLE EXPERTS NAMES)',         author_line)
    _replace_global(body, '(RESPONSIBLE EXPERTS SIGNATURES + INK PAD)', name_line)
    _replace_global(body, '(RESPONSIBLE EXPERTS POSITIONS)',     reviewer_line)

    # Ref. body paragraph
    for p in body.iter(f'{{{W}}}p'):
        if _get_text(p).strip() == 'Ref.:':
            _set_para_text(p, 'Ref.:  ' + data.get('short_ref', ''))
            break

    # Tawuniya Visit ID
    for p in body.iter(f'{{{W}}}p'):
        if _get_text(p).strip() == 'Tawuniya visit ID:':
            _set_para_text(p, 'TAWUNIYA Visit ID:  ' + data.get('taw_visit_id', ''))
            break


def _insert_engineer_signature(doc, eng_full, sig_bytes):
    """Insert engineer's signature image after each name/reviewer paragraph."""
    if not sig_bytes:
        return
    from io import BytesIO
    target = 'Eng. ' + eng_full
    for para in doc.paragraphs:
        if target in para.text and 'Reviewer:' in para.text:
            try:
                run = para.add_run()
                from docx.oxml import OxmlElement as OXE
                br = OXE('w:br')
                run._r.append(br)
                run.add_picture(BytesIO(sig_bytes), width=Cm(2.5))
            except Exception:
                pass


# ── Table 1 — Document header + expert info ───────────────────────────────────

def _fill_table1(table, data):
    """Fill T1: doc ref, date, expert name/phase/degree/speciality, author block."""
    rd3_ref    = data.get('rd3_ref', '')
    issue_date = data.get('issue_date', '')
    eng_full   = data.get('eng_full', '')
    eng_phase  = data.get('eng_phase', 'Waterproofing')
    eng_degree = data.get('eng_degree', 'Bachelor')
    eng_spec   = data.get('eng_speciality', 'Civil Engineer')
    eng_phone  = data.get('eng_phone', '')
    eng_email  = data.get('eng_email', '')

    # Row 0: DOCUMENT REFERENCE & TIS AGENCY
    _append_to_label_para(table.rows[0].cells[0]._tc, 'DOCUMENT REFERENCE', rd3_ref)
    _append_to_label_para(table.rows[0].cells[2]._tc, 'TIS AGENCY',         'SOCOTEC ARABIA')
    # Row 1: VERSION & DATE OF ISSUE
    _append_to_label_para(table.rows[1].cells[0]._tc, 'VERSION',            '1')
    _append_to_label_para(table.rows[1].cells[2]._tc, 'DATE OF ISSUE',      issue_date)

    # Row 4: Expert details — use python-docx API (cell[0] has 5 empty paragraphs)
    _fill_expert_cell(table, 4, 0, eng_full)
    _fill_expert_cell(table, 4, 1, eng_phase)
    _fill_expert_cell(table, 4, 2, eng_degree)
    _fill_expert_cell(table, 4, 3, eng_spec)

    # Row 5: AUTOR block — use python-docx API for styled output
    cell5 = table.rows[5].cells[0]
    for para in cell5.paragraphs:
        txt = para.text.strip()
        if txt.startswith('AUTOR of THIS REPORT'):
            para.clear()
            r1 = para.add_run('AUTOR of THIS REPORT: ')
            r1.bold = True; r1.font.size = Pt(9)
            r2 = para.add_run(eng_full)
            r2.bold = True; r2.font.size = Pt(9); r2.font.color.rgb = BLUE
        elif txt.startswith('PHONE NUMBER'):
            para.clear()
            r1 = para.add_run('PHONE NUMBER: ')
            r1.bold = True; r1.font.size = Pt(9)
            r2 = para.add_run(eng_phone)
            r2.bold = True; r2.font.size = Pt(9); r2.font.color.rgb = BLUE
        elif txt.startswith('EMAIL'):
            para.clear()
            r1 = para.add_run('EMAIL: ')
            r1.bold = True; r1.font.size = Pt(9)
            r2 = para.add_run(eng_email)
            r2.bold = True; r2.font.size = Pt(9); r2.font.color.rgb = BLUE


# ── Table 2 — Main report body ────────────────────────────────────────────────

def _fill_table2(table, data, visits):
    """Fill T2: project info, occupancy, visits, conclusion."""
    # Row 1: project info cell
    proj_cell = table.rows[1].cells[0]
    info_map = {
        'PROJECT TITLE (NAME)':                      data.get('project_title', ''),
        'ADDRESS OF THE PREMISES':                   data.get('address', ''),
        'REFERENCE RD0':                             data.get('rd0_ref', ''),
        'PRINCIPAL/OWNER':                           data.get('owner', ''),
        'BUILDINGS INCLUDED IN THE PROJECT AND ITS USE': data.get('buildings', '1 residential building'),
    }
    for para in proj_cell.paragraphs:
        txt = para.text.strip()
        for label, val in info_map.items():
            if txt.startswith(label) and val:
                para.clear()
                r1 = para.add_run(txt.rstrip() + ' ')
                r1.bold = True; r1.font.size = Pt(9)
                r2 = para.add_run(str(val))
                r2.bold = True; r2.font.size = Pt(9); r2.font.color.rgb = BLUE
                break

    # Row 5 col 2: Occupancy date + Expected/Confirmed
    occ_cell = table.rows[5].cells[2]
    for para in occ_cell.paragraphs:
        txt = para.text
        if 'Date of Occupancy Certificate' in txt:
            para.clear()
            r1 = para.add_run('Date of Occupancy Certificate: ')
            r1.bold = True; r1.font.size = Pt(9)
            r2 = para.add_run(data.get('occ_date', ''))
            r2.bold = True; r2.font.size = Pt(9); r2.font.color.rgb = BLUE
        elif '\u2610 Expected' in txt and '\u2610 Confirmed' in txt:
            occ = data.get('occ_status', 'Expected')
            if occ == 'Expected':
                new = txt.replace('\u2610 Expected', '\u2612 Expected')
            else:
                new = txt.replace('\u2610 Confirmed', '\u2612 Confirmed')
            _set_para_text(para._p, new)

    # Row 6 col 2: site visits nested table
    visit_cell = table.rows[6].cells[2]._tc
    nested_tbls = visit_cell.findall(f'.//{{{W}}}tbl')
    if nested_tbls and visits:
        rows = nested_tbls[0].findall(f'{{{W}}}tr')
        for i, v in enumerate(visits[:10]):
            row_idx = i + 1
            if row_idx >= len(rows):
                break
            cells = rows[row_idx].findall(f'.//{{{W}}}tc')
            for tc, val in zip(cells, [v.get('ref',''), v.get('date',''),
                                        v.get('inspector',''), v.get('part','')]):
                _fill_cell_text_tc(tc, val)

    # Row 11: conclusion — use _tc.iter to reach nested YES/NO paragraph
    if len(table.rows) > 11:
        conc_tc = table.rows[11].cells[0]._tc
        conc_yn = data.get('conclusion_yn', 'YES')
        for p in conc_tc.iter(f'{{{W}}}p'):
            ptxt = _get_text(p)
            if '☐ YES' in ptxt and '☐ NO' in ptxt:
                new_txt = ptxt.replace('☐ YES', '☒ YES', 1) if conc_yn == 'YES' else \
                          ptxt.replace('☐ NO',  '☒ NO',  1)
                _set_para_text(p, new_txt)
                break
        conc_text = data.get('conclusion_text', '')
        if conc_text:
            _append_text_to_cell_last_para(conc_tc, conc_text)


# ── Annex table filler ────────────────────────────────────────────────────────

def _fill_annex_table(table, ann, annex_type):
    """Fill annex table (T3=roofs, T4=facades, T5=basements)."""
    # Row 2: Works description
    if len(table.rows) > 2:
        desc_cell = table.rows[2].cells[0]
        for para in desc_cell.paragraphs:
            txt = para.text.strip()
            if txt.startswith('I.1.') and ann.get('description_i1'):
                para.clear()
                r1 = para.add_run(txt + '\n\n')
                r1.font.size = Pt(9)
                r2 = para.add_run(ann['description_i1'])
                r2.font.size = Pt(9); r2.font.color.rgb = BLUE
            elif txt.startswith('I.2.') and ann.get('description_i2'):
                para.clear()
                r1 = para.add_run(txt + '\n\n')
                r1.font.size = Pt(9)
                r2 = para.add_run(ann['description_i2'])
                r2.font.size = Pt(9); r2.font.color.rgb = BLUE
            elif txt.startswith('I.3.') and ann.get('description_i3'):
                para.clear()
                r1 = para.add_run(txt + '\n\n')
                r1.font.size = Pt(9)
                r2 = para.add_run(ann['description_i3'])
                r2.font.size = Pt(9); r2.font.color.rgb = BLUE
            elif 'I.4.' in txt and '☐' in txt:
                # I.4 innovative technique YES/NO — ALL annex types
                inn = ann.get('innovative_yn', 'NO')
                para.clear()
                r1 = para.add_run(txt.replace('☐ YES', '☒ YES' if inn == 'YES' else '☐ YES')
                                     .replace('☐ NO',  '☒ NO'  if inn == 'NO'  else '☐ NO'))
                r1.font.size = Pt(9)

    # Materials row
    mat_row = 4 if annex_type != 'basement' else 3
    if len(table.rows) > mat_row:
        _fill_materials_cell(table.rows[mat_row].cells[0], ann)

    # Conclusion — use _tc.iter to reach nested YES/NO paragraph
    last_tc = table.rows[-1].cells[0]._tc
    conc_yn = ann.get('conclusion_yn', 'YES')
    for p in last_tc.iter(f'{{{W}}}p'):
        txt = _get_text(p)
        if '☐ YES' in txt and '☐ NO' in txt:
            new_txt = txt.replace('☐ YES', '☒ YES') if conc_yn == 'YES' else \
                      txt.replace('☐ NO',  '☒ NO')
            _set_para_text(p, new_txt)
        elif 'Please expose' in txt or 'developpe' in txt.lower():
            if ann.get('conclusion_text'):
                r = etree.SubElement(p, f'{{{W}}}r')
                te = etree.SubElement(r, f'{{{W}}}t')
                te.text = '\n\n' + ann['conclusion_text']
                te.set(XML_SPACE, 'preserve')


def _fill_materials_cell(cell, ann):
    """Fill reviewed documents lists in the materials cell."""
    docs_mat  = ann.get('reviewed_docs_materials', [])
    docs_test = ann.get('reviewed_docs_tests', [])
    docs_other= ann.get('reviewed_docs_other', [])

    current_section = None
    doc_counter = {'mat': 0, 'test': 0, 'other': 0}

    for para in cell.paragraphs:
        txt = para.text.strip()
        if 'Delivery Orders - Material certificates' in txt:
            current_section = 'mat'
        elif 'Tests and quality control reports' in txt:
            current_section = 'test'
        elif 'Oher Reviewed documents' in txt or 'Other Reviewed' in txt:
            current_section = 'other'
        elif txt and txt[0].isdigit() and '.' in txt[:3]:
            num = txt.split('.')[0]
            docs = docs_mat if current_section == 'mat' else (
                   docs_test if current_section == 'test' else docs_other)
            key = current_section or 'mat'
            idx = doc_counter.get(key, 0)
            if idx < len(docs) and docs[idx]:
                para.clear()
                r = para.add_run('\t{}. {}'.format(num, docs[idx]))
                r.font.size = Pt(9); r.font.color.rgb = BLUE
                doc_counter[key] = idx + 1


# ── SDT checkbox logic ────────────────────────────────────────────────────────

def _toggle_all_sdts(body, annex_data):
    has_roofs     = 'roofs'     in annex_data
    has_facades   = 'facades'   in annex_data
    has_basements = 'basements' in annex_data

    _toggle_sdt(body, 0, has_roofs)
    _toggle_sdt(body, 1, has_facades)
    _toggle_sdt(body, 2, has_basements)

    if has_roofs:
        rt = annex_data['roofs'].get('roof_types', [])
        _toggle_sdt(body, 3, 'ROOF'                 in rt)
        _toggle_sdt(body, 4, 'ROOFTOP TERRACE'      in rt)
        _toggle_sdt(body, 5, 'INTERMEDIATE TERRACE' in rt)
        _toggle_sdt(body, 6, 'PATIOS'               in rt)
        _toggle_sdt(body, 7, 'OTHER'                in rt)
        rd = annex_data['roofs']
        _toggle_sdt(body, 8,  rd.get('delivery_yn',  'YES') == 'YES')
        _toggle_sdt(body, 9,  rd.get('delivery_yn',  'YES') == 'NO')
        _toggle_sdt(body, 10, rd.get('compliant_yn', 'YES') == 'YES')
        _toggle_sdt(body, 11, rd.get('compliant_yn', 'YES') == 'NO')
        _toggle_sdt(body, 12, rd.get('ponding_yn',   'YES') == 'YES')
        _toggle_sdt(body, 13, rd.get('ponding_yn',   'YES') == 'NO')

    if has_facades:
        ft = annex_data['facades'].get('facade_types', [])
        _toggle_sdt(body, 14, 'CONCRETE OR MASONRY' in ft)
        _toggle_sdt(body, 15, 'CLADDING'            in ft)
        _toggle_sdt(body, 16, 'CURTAIN WALL'        in ft)
        _toggle_sdt(body, 17, 'OTHER'               in ft)
        fd = annex_data['facades']
        _toggle_sdt(body, 18, fd.get('delivery_yn',  'YES') == 'YES')
        _toggle_sdt(body, 19, fd.get('delivery_yn',  'YES') == 'NO')
        _toggle_sdt(body, 20, fd.get('compliant_yn', 'YES') == 'YES')
        _toggle_sdt(body, 21, fd.get('compliant_yn', 'YES') == 'NO')
        _toggle_sdt(body, 22, fd.get('ponding_yn',   'YES') == 'YES')
        _toggle_sdt(body, 23, fd.get('ponding_yn',   'YES') == 'NO')

    if has_basements:
        bd = annex_data['basements']
        _toggle_sdt(body, 24, bd.get('delivery_yn',  'YES') == 'YES')
        _toggle_sdt(body, 25, bd.get('delivery_yn',  'YES') == 'NO')
        _toggle_sdt(body, 26, bd.get('compliant_yn', 'YES') == 'YES')
        _toggle_sdt(body, 27, bd.get('compliant_yn', 'YES') == 'NO')
        _toggle_sdt(body, 28, bd.get('ponding_yn',   'YES') == 'YES')
        _toggle_sdt(body, 29, bd.get('ponding_yn',   'YES') == 'NO')


# ── Attachments ───────────────────────────────────────────────────────────────

def _add_attachments(doc, attachments):
    """Add attachments at end of document. Images displayed inline; PDFs as styled reference."""
    if not attachments:
        return
    from io import BytesIO
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    # Header
    para = doc.add_paragraph()
    run = para.add_run('ATTACHMENTS')
    run.bold = True; run.font.size = Pt(12); run.font.color.rgb = BLUE

    pdf_refs = []  # collect PDFs to list at end

    for fname, fbytes in attachments:
        ext = fname.lower().split('.')[-1] if '.' in fname else ''

        if ext == 'pdf':
            # Don't add a page break for PDFs — just collect for reference list
            pdf_refs.append(fname)
        else:
            # Page break before each image attachment
            bp = doc.add_paragraph()
            r = OxmlElement('w:r'); br = OxmlElement('w:br')
            br.set(qn('w:type'), 'page'); r.append(br); bp._p.append(r)
            # Caption
            cap = doc.add_paragraph()
            cap.add_run(fname).bold = True
            # Image
            try:
                img_para = doc.add_paragraph()
                img_para.add_run().add_picture(BytesIO(fbytes), width=Cm(15))
            except Exception:
                doc.add_paragraph('[ Image could not be rendered: {} ]'.format(fname))

    # PDF reference section — clean list at the end
    if pdf_refs:
        doc.add_paragraph()
        ref_head = doc.add_paragraph()
        rh = ref_head.add_run('Documents to Attach (print and include with this report):')
        rh.bold = True; rh.font.size = Pt(10)
        for i, fname in enumerate(pdf_refs, 1):
            p = doc.add_paragraph()
            r = p.add_run('{}. {}'.format(i, fname))
            r.font.size = Pt(10); r.font.color.rgb = BLUE


# ── Break type fix ────────────────────────────────────────────────────────────

def _add_page_break(doc):
    """Add a page break paragraph."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    para = doc.add_paragraph()
    r = OxmlElement('w:r')
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r.append(br)
    para._p.append(r)
    return para


def _fill_annex_other_text(table, other_text):
    """
    Fill the OTHER: text in annex type row (row 0).
    Paragraph structure: [tab][tab][SDT ☐][run 'OTHER: '][run '\t']
    Must modify only the 'OTHER: ' run — NOT use _set_para_text (which destroys the SDT).
    """
    if not other_text:
        return
    type_cell = table.rows[0].cells[0]._tc
    for p in type_cell.findall(f'.//{{{W}}}p'):
        # Only look at direct child w:r runs (skips SDT content)
        for r in p.findall(f'{{{W}}}r'):
            for t_elem in r.findall(f'{{{W}}}t'):
                if t_elem.text and 'OTHER:' in t_elem.text:
                    # Replace only the text of this run
                    t_elem.text = 'OTHER: ' + other_text
                    t_elem.set(XML_SPACE, 'preserve')
                    return


def _enrich_reviewed_docs(ann, visit_refs):
    """
    Auto-fill reviewed docs if engineer left them blank.
    Smart defaults: warranty certificate + visit report references.
    """
    # Material docs: warranty certificate if empty
    if not ann.get('reviewed_docs_materials'):
        ann['reviewed_docs_materials'] = ['Waterproofing warranty certificate']

    # Test docs: warranty + visit report refs if empty
    if not ann.get('reviewed_docs_tests'):
        test_docs = ['Waterproofing warranty certificate']
        # Add intermediate visit references (all except the very last = final visit)
        for ref in visit_refs[:-1]:   # skip last visit (final)
            if ref:
                test_docs.append('Visit Report ({})'.format(ref))
        # If only 1 visit total, still add it as test doc
        if len(visit_refs) == 1 and visit_refs[0]:
            test_docs.append('Visit Report ({})'.format(visit_refs[0]))
        ann['reviewed_docs_tests'] = test_docs[:4]  # max 4 slots


def _mark_annex_na(table):
    """Mark a non-selected annex table as NOT APPLICABLE — fills all empty text areas with N/A."""
    na_run_added = set()
    for ri, row in enumerate(table.rows):
        cell = row.cells[0]
        txt = cell.text.strip()
        # Fill description cells (I.1, I.2, I.3) and conclusion text with N/A
        if any(txt.startswith(lbl) for lbl in ('I.1.', 'I.2.', 'I.3.')):
            for para in cell.paragraphs:
                if not para.text.strip():
                    run = para.add_run('N/A')
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
                    break


# ── Main entry point ──────────────────────────────────────────────────────────

def generate_rd3(template_path, output_path, data, visits, annex_data,
                 sig_bytes=None, attachments=None):
    """
    Fill Template_RD3.docx and save to output_path.

    Parameters
    ----------
    template_path : str
    output_path   : str
    data          : dict with engineer/project/report fields
    visits        : list of {ref, date, inspector, part}
    annex_data    : dict — keys 'roofs', 'facades', 'basements'
    sig_bytes     : bytes of engineer signature PNG (optional)
    attachments   : list of (filename, bytes) tuples (optional)
    """
    doc  = Document(template_path)
    body = doc.element.body

    # 1. SDT checkboxes
    _toggle_all_sdts(body, annex_data)

    # 2. Signature blocks (body paragraphs)
    _fill_signature_blocks(body, data)

    # 3. Table 1: header + expert
    _fill_table1(doc.tables[1], data)

    # 4. Table 2: main report
    _fill_table2(doc.tables[2], data, visits)
    _fill_t2_defects_reservations(doc.tables[2], data)

    # 5. Annex tables
    # Enrich annex data with visit refs for reviewed docs auto-fill
    visit_refs = [v.get('ref','') for v in visits if v.get('ref','')]

    if 'roofs' in annex_data and len(doc.tables) > 3:
        _fill_annex_other_text(doc.tables[3], annex_data['roofs'].get('other_type_text',''))
        _enrich_reviewed_docs(annex_data['roofs'], visit_refs)
        _fill_annex_table(doc.tables[3], annex_data['roofs'], 'roof')
        _fill_annex_extras(doc.tables[3], annex_data['roofs'], 'roof')
    elif len(doc.tables) > 3:
        _mark_annex_na(doc.tables[3])

    if 'facades' in annex_data and len(doc.tables) > 4:
        _fill_annex_other_text(doc.tables[4], annex_data['facades'].get('other_type_text',''))
        _enrich_reviewed_docs(annex_data['facades'], visit_refs)
        _fill_annex_table(doc.tables[4], annex_data['facades'], 'facade')
        _fill_annex_extras(doc.tables[4], annex_data['facades'], 'facade')
    elif len(doc.tables) > 4:
        _mark_annex_na(doc.tables[4])

    if 'basements' in annex_data and len(doc.tables) > 5:
        _enrich_reviewed_docs(annex_data['basements'], visit_refs)
        _fill_annex_table(doc.tables[5], annex_data['basements'], 'basement')
        _fill_annex_extras(doc.tables[5], annex_data['basements'], 'basement')
    elif len(doc.tables) > 5:
        _mark_annex_na(doc.tables[5])

    # 6. Engineer signature image (inserted into each signature block)
    if sig_bytes:
        _insert_engineer_signature(doc, data.get('eng_full', ''), sig_bytes)

    # 7. Attachments at end
    if attachments:
        _add_attachments(doc, attachments)

    doc.save(output_path)
    return output_path


# ═══ ADDITIONAL HELPERS FOR MISSING SECTIONS ═══════════════════════════════


def _toggle_adjacent_cb(cell_tc, nested_tbl_idx, row_idx, answer):
    """Toggle ☐ in cell1 of a nested table row (question in cell0, checkbox in cell1)."""
    ntbls = cell_tc.findall(f'.//{{{W}}}tbl')
    if nested_tbl_idx >= len(ntbls): return False
    tr_list = ntbls[nested_tbl_idx].findall(f'{{{W}}}tr')
    if row_idx >= len(tr_list): return False
    for tc in reversed(tr_list[row_idx].findall(f'.//{{{W}}}tc')):
        txt = ''.join(t.text or '' for t in tc.iter(f'{{{W}}}t'))
        if '☐' in txt:
            for p in tc.iter(f'{{{W}}}p'):
                ptxt = ''.join(t.text or '' for t in p.iter(f'{{{W}}}t'))
                if '☐' in ptxt:
                    new = ptxt.replace('☐ YES', '☒ YES', 1) if answer == 'YES' else \
                          ptxt.replace('☐ NO',  '☒ NO',  1)
                    _set_para_text(p, new)
                    return True
    return False


def _toggle_standalone_cb_by_phrase(cell_tc, phrase, answer):
    """Find a row containing phrase and toggle its adjacent checkbox."""
    ntbls = cell_tc.findall(f'.//{{{W}}}tbl')
    for ni, nt in enumerate(ntbls):
        for ri, tr in enumerate(nt.findall(f'{{{W}}}tr')):
            row_txt = ''.join(t.text or '' for t in tr.iter(f'{{{W}}}t'))
            if phrase.lower() in row_txt.lower():
                if _toggle_adjacent_cb(cell_tc, ni, ri, answer):
                    return True
    # Fallback: search direct paragraphs
    for p in cell_tc.iter(f'{{{W}}}p'):
        ptxt = ''.join(t.text or '' for t in p.iter(f'{{{W}}}t'))
        if phrase.lower() in ptxt.lower() and '☐' in ptxt:
            new = ptxt.replace('☐ YES', '☒ YES', 1) if answer == 'YES' else \
                  ptxt.replace('☐ NO',  '☒ NO',  1)
            _set_para_text(p, new)
            return True
    return False


def _fill_annex_extras(tbl, ann, annex_type):
    """Fill modifications, defects, reservations, and conclusion checkboxes in annex tables."""
    mod_row, def_row, res_row = (5, 7, 9) if annex_type == 'basement' else (6, 8, 10)
    mods_yn = ann.get('modifications_yn', 'NO')
    def_yn  = ann.get('defects_observed_yn', 'NO')

    # ── III Modifications ────────────────────────────────────────────────────
    if mod_row < len(tbl.rows):
        mod_tc = tbl.rows[mod_row].cells[0]._tc
        _toggle_adjacent_cb(mod_tc, 0, 0, mods_yn)   # main question: nested tbl 0, row 0
        if mods_yn == 'YES':
            _toggle_adjacent_cb(mod_tc, 1, 0, ann.get('mod_validation_yn', 'YES'))
            _toggle_adjacent_cb(mod_tc, 1, 2, ann.get('mod_acceptable_yn', 'YES'))
            _toggle_adjacent_cb(mod_tc, 1, 4, ann.get('mod_execution_yn',  'YES'))

    # ── IV Defects / Waterproofing Tests ─────────────────────────────────────
    if def_row < len(tbl.rows):
        def_tc = tbl.rows[def_row].cells[0]._tc
        _toggle_adjacent_cb(def_tc, 0, 0, def_yn)    # defect question: nested tbl 0, row 0
        if def_yn == 'YES':
            ref = ann.get('defects_ref', '')
            if ref:
                for p in def_tc.iter(f'{{{W}}}p'):
                    ptxt = ''.join(t.text or '' for t in p.iter(f'{{{W}}}t'))
                    if 'Ref.:' in ptxt and 'please issue' in ptxt.lower():
                        t_elems = list(p.iter(f'{{{W}}}t'))
                        if t_elems: t_elems[-1].text = (t_elems[-1].text or '') + ' ' + ref
                        break

    # ── V Reservations ────────────────────────────────────────────────────────
    if res_row < len(tbl.rows):
        res_tc = tbl.rows[res_row].cells[0]._tc
        _toggle_standalone_cb_by_phrase(res_tc, 'issued and closed',
                                        ann.get('reservations_closed_yn', 'NO'))
        _toggle_standalone_cb_by_phrase(res_tc, 'issued not closed',
                                        ann.get('reservations_open_yn',   'NO'))

    # ── Conclusion YES/NO — search all paragraphs in last row via _tc.iter ────
    last_tc = tbl.rows[-1].cells[0]._tc
    conc_yn = ann.get('conclusion_yn', 'YES')
    for p in last_tc.iter(f'{{{W}}}p'):
        ptxt = ''.join(t.text or '' for t in p.iter(f'{{{W}}}t'))
        if ptxt.strip() in ('☐ YES☐ NO', '☐ YES ☐ NO') or \
           (ptxt.strip().startswith('☐') and 'YES' in ptxt and 'NO' in ptxt):
            new = ptxt.replace('☐ YES', '☒ YES', 1) if conc_yn == 'YES' else \
                  ptxt.replace('☐ NO',  '☒ NO',  1)
            _set_para_text(p, new)
            break


def _fill_t2_defects_reservations(table, data):
    """Fill T2 row 7 (defects) and row 9 (reservations checkboxes)."""
    defects = data.get('defects_text', 'None')

    # Row 7: defects text
    if len(table.rows) > 7:
        cell7 = table.rows[7].cells[0]
        paras = cell7.paragraphs
        for pi, para in enumerate(paras):
            if 'Please list below' in para.text:
                for ep in paras[pi+1:]:
                    if not ep.text.strip():
                        run = ep.add_run(defects)
                        run.font.size = Pt(9); run.font.color.rgb = BLUE
                        break
                break

    # Row 9: toggle reservations checkboxes (do NOT add "None" text above the table)
    if len(table.rows) > 9:
        cell9_tc = table.rows[9].cells[0]._tc
        has_res = data.get('reservations_text','').strip().lower() not in ('none','n/a','')
        _toggle_standalone_cb_by_phrase(cell9_tc, 'issued not closed', 'YES' if has_res else 'NO')
        _toggle_standalone_cb_by_phrase(cell9_tc, 'issued and closed', 'NO')
